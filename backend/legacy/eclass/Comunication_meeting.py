from typing import Dict, List, Tuple, Optional
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from typing import Dict, Optional
import numpy as np
import re  # 新增：正则解析日期/成绩
from datetime import datetime  # 新增：日期处理
from docx.oxml.ns import qn
import os
import warnings
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import pandas as pd


# 需要作为“分公司”维度单独统计的特殊组织名称。
# 这些组织不含“分公司”字样，原逻辑容易解析为空或归到上一级。
SPECIAL_BRANCH_NAMES = [
    "深圳办事处",
    "体外诊断临床应用部",
    "北京战略客户服务部",
    "临检用户服务部",
    "中国区体外诊断临床应用部",
]

UNKNOWN_BRANCH_NAME = "未知分公司"

# 以下分公司仍保留在名单/成绩明细中，但不参与附录分公司统计和 Word 总结/图表。
EXCLUDED_WORD_AND_BRANCH_SUMMARY_BRANCHES = [
    "临检用户服务部",
    "体外诊断临床应用部",
]


def load_attendance_roster(file_path: str) -> pd.DataFrame:
    """加载应出勤人员基准名单"""
    try:
        df = pd.read_excel(
            file_path,
            sheet_name='Sheet1',
            dtype={'工号': str},
            usecols=['工号', '姓名', '分公司']
        ).dropna(how='all')

        df['工号'] = df['工号'].str.strip().str.upper()
        df = df.drop_duplicates(subset=['工号'], keep='first')
        df['工号'] = df['工号'].fillna('UNKNOWN')
        df['分公司'] = df['分公司'].str.strip().str.replace(r'\s+', ' ', regex=True)

        assert df['工号'].nunique() == len(df), "工号存在重复值，请检查名单.xlsx"
        return df.reset_index(drop=True)

    except Exception as e:
        raise ValueError(f"加载应出勤名单失败: {str(e)}") from e

def fill_unattended_subjects(row, exam_types: list):
    """
    将未参加的科目填充为"/"，参加的科目保留原分数（包括0分）
    :param row: DataFrame行数据（需包含attended_subjects字段）
    :param exam_types: 所有考试科目列表
    :return: 处理后的科目成绩字典
    """
    attended_subjects = row.get('attended_subjects', [])
    result = {}
    for subject in exam_types:
        if subject in attended_subjects:
            # 参加过的科目保留原分数（包括0分）
            result[subject] = row[subject] if pd.notna(row[subject]) else 0
        else:
            # 未参加的科目填充为"/"
            result[subject] = "/"
    return result

# 【其余函数保持不变，仅修改 load_exam_scores 函数】

def _find_header_row(file_path: str, required_cols=None, max_scan_rows: int = 30) -> int:
    """
    自动识别 Excel 表头所在行。

    新版考试导出表前面可能有多行标题，例如：
    2026年05月凝血技术交流会考试
    2026年05月凝血技术交流会考试
    ...
    账号 姓名 考试状态 部门 开考时间 交卷时间 成绩 ...

    所以不能直接 pd.read_excel(header=0)，需要先扫描表头行。
    """
    if required_cols is None:
        required_cols = ["账号", "部门", "成绩"]

    preview = pd.read_excel(
        file_path,
        header=None,
        nrows=max_scan_rows,
        dtype=str,
        engine="openpyxl"
    )

    for idx, row in preview.iterrows():
        values = [str(x).strip() for x in row.tolist() if pd.notna(x)]
        if all(col in values for col in required_cols):
            return idx

    raise ValueError(
        f"未找到有效表头行，要求至少包含列：{required_cols}，文件：{file_path}"
    )


def _extract_dispatch_date_from_filename(file_path: str):
    """
    从文件名中提取考试派发日期。

    支持格式：
    2026-05-28考试考生情况-2026年5月中国区技术交流会考试--化免20260608114705.xlsx
    2026-05-29考生情况-2026年05月凝血技术交流会考试20260608114558.xlsx
    2026-04-27-临检成绩.xlsx
    """
    filename = os.path.basename(file_path)

    # 优先提取文件名前面的标准日期：2026-05-28
    match = re.search(r"(\d{4})[-_/\.](\d{1,2})[-_/\.](\d{1,2})", filename)
    if match:
        year, month, day = match.groups()
        return pd.to_datetime(f"{int(year):04d}-{int(month):02d}-{int(day):02d}").date()

    # 兼容中文日期：2026年5月28日
    match = re.search(r"(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日", filename)
    if match:
        year, month, day = match.groups()
        return pd.to_datetime(f"{int(year):04d}-{int(month):02d}-{int(day):02d}").date()

    raise ValueError(
        f"无法从文件名中提取考试派发日期，请确认文件名包含 YYYY-MM-DD 或 YYYY年M月D日 格式日期：{filename}"
    )


def _extract_branch_from_department(department):
    """
    从部门路径中提取分公司。

    示例：
    签约渠道商(外部)->杭州分公司->宁波英特天泽鼎丰贸易有限公司
    提取：杭州分公司

    全球用户服务部->中国区用户服务部->西南区用户服务部->成都分公司
    提取：成都分公司
    """
    if pd.isna(department):
        return ""

    text = str(department).strip()
    if not text:
        return ""

    # 兼容不同箭头写法
    text = text.replace("－>", "->").replace("—>", "->").replace("→", "->")
    parts = [p.strip() for p in text.split("->") if p and p.strip()]

    # 优先找包含“分公司”的层级
    for part in parts:
        if "分公司" in part:
            return part

    # 如果没有“分公司”，兼容“办事处”“服务部”等情况
    for part in reversed(parts):
        if any(keyword in part for keyword in ["办事处", "服务部", "战略客户服务部"]):
            return part

    # 兜底：取最后一级
    return parts[-1] if parts else ""


def _to_numeric_score(value) -> float:
    """
    将成绩转换为数值。

    兼容：
    100
    100.0
    '100分'
    '--'
    空值
    """
    if pd.isna(value):
        return 0.0

    text = str(value).strip()

    if text in ["", "--", "-", "/", "未参考", "未提交", "nan", "None"]:
        return 0.0

    # 优先直接转数字
    try:
        return float(text)
    except Exception:
        pass

    # 从字符串中提取最后一个数字，避免“第1次考试 85分”误取 1
    nums = re.findall(r"\d+(?:\.\d+)?", text)
    if nums:
        return float(nums[-1])

    return 0.0


def _normalize_account(value) -> str:
    """
    统一账号/工号格式。
    避免 Excel 将纯数字账号读成 50259563.0。
    """
    if pd.isna(value):
        return ""

    text = str(value).strip()

    if text.endswith(".0"):
        text = text[:-2]

    return text


def _calculate_adjusted_score(raw_score, submit_time, dispatch_date, exam_status=None):
    """
    根据考试派发日期修正成绩。

    规则：
    1. 派发当天考试：最终成绩 = 原始成绩
    2. 派发日期之后考试：
       - 原始成绩 >= 80，则最终成绩 = 80
       - 原始成绩 < 80，则最终成绩 = 原始成绩
    3. 未开始/未提交/无交卷时间：最终成绩 = 0
    """
    raw_score = _to_numeric_score(raw_score)

    status_text = "" if pd.isna(exam_status) else str(exam_status).strip()

    # 未开始、未提交，直接按 0
    if status_text in ["未开始", "未提交"]:
        return 0.0

    submit_dt = pd.to_datetime(submit_time, errors="coerce")

    # 没有有效交卷时间，视为未参加
    if pd.isna(submit_dt):
        return 0.0

    submit_date = submit_dt.date()

    # 派发当天考试，保留原始成绩
    if submit_date == dispatch_date:
        return raw_score

    # 派发日期之后考试，最高按 80
    if submit_date > dispatch_date:
        return min(raw_score, 80.0) if raw_score >= 80 else raw_score

    # 理论上不会早于派发日期，如出现则保留原始成绩
    return raw_score

def _find_header_row(file_path: str, required_cols=None, max_scan_rows: int = 40) -> int:
    """
    自动识别 Excel 表头行。
    新版成绩表前面有多行考试标题，真正表头是：
    账号、姓名、考试状态、部门、交卷时间、成绩...
    """
    if required_cols is None:
        required_cols = ["账号", "部门", "成绩"]

    preview = pd.read_excel(
        file_path,
        header=None,
        nrows=max_scan_rows,
        dtype=str,
        engine="openpyxl"
    )

    for idx, row in preview.iterrows():
        values = [str(x).strip() for x in row.tolist() if pd.notna(x)]
        if all(col in values for col in required_cols):
            return idx

    raise ValueError(
        f"未找到有效表头行，要求至少包含列：{required_cols}，文件：{file_path}"
    )


def _extract_dispatch_date_from_filename(file_path: str):
    """
    从文件名中提取考试派发日期。

    支持：
    2026-05-28考试考生情况-2026年5月中国区技术交流会考试--化免20260608114705.xlsx
    2026-05-29考生情况-2026年05月凝血技术交流会考试20260608114558.xlsx
    2026-04-27-临检成绩.xlsx
    """
    filename = os.path.basename(file_path)

    match = re.search(r"(\d{4})-(\d{1,2})-(\d{1,2})", filename)
    if match:
        year, month, day = match.groups()
        return pd.to_datetime(f"{int(year):04d}-{int(month):02d}-{int(day):02d}").date()

    match = re.search(r"(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日", filename)
    if match:
        year, month, day = match.groups()
        return pd.to_datetime(f"{int(year):04d}-{int(month):02d}-{int(day):02d}").date()

    raise ValueError(
        f"无法从文件名中提取考试派发日期，请确认文件名包含 YYYY-MM-DD 或 YYYY年M月D日：{filename}"
    )


def _normalize_account(value) -> str:
    """
    统一账号格式，避免 50259563.0 这种情况。
    """
    if pd.isna(value):
        return ""

    text = str(value).strip()

    if text.endswith(".0"):
        text = text[:-2]

    return text


def _to_numeric_score(value) -> float:
    """
    成绩转数字。
    """
    if pd.isna(value):
        return 0.0

    text = str(value).strip()

    if text in ["", "--", "-", "/", "未参考", "未提交", "nan", "None"]:
        return 0.0

    try:
        return float(text)
    except Exception:
        pass

    nums = re.findall(r"\d+(?:\.\d+)?", text)
    if nums:
        return float(nums[-1])

    return 0.0


def _parse_department_info(department):
    """
    从部门字段解析分公司和渠道商。

    兼容：
    1. 普通分公司：上海分公司、成都分公司等；
    2. 特殊内部组织：深圳办事处、体外诊断临床应用部、北京战略客户服务部、临检用户服务部、中国区体外诊断临床应用部；
    3. 渠道：签约渠道商(外部)->上海分公司->上海其田医疗科技有限公司。
    """
    if pd.isna(department):
        return "", ""

    text = str(department).strip()

    if not text:
        return "", ""

    text = (
        text.replace("&gt;", ">")
            .replace("－>", "->")
            .replace("—>", "->")
            .replace("→", "->")
            .replace("＞", ">")
    )

    # 兼容 ->、- >、—>、=> 等箭头
    parts = re.split(r"\s*[-－—=]*>\s*", text)
    parts = [p.strip() for p in parts if p and p.strip()]

    branch = ""
    channel = ""

    # 优先识别真正的“分公司”层级
    for part in parts:
        if "分公司" in part:
            branch = part
            break

    # 再识别需要额外统计的特殊组织，避免这些内部员工分公司为空
    if not branch:
        for part in parts:
            for special_branch in SPECIAL_BRANCH_NAMES:
                if special_branch == part or special_branch in part:
                    branch = special_branch
                    break
            if branch:
                break

    is_channel = "签约渠道商" in text

    if is_channel:
        if branch and branch in parts:
            idx = parts.index(branch)
            if idx + 1 < len(parts):
                channel = parts[idx + 1]

        if not channel and len(parts) >= 3:
            channel = parts[-1]

    return branch, channel

def _is_blank_value(value) -> bool:
    """
    判断单元格是否为空值。
    """
    if pd.isna(value):
        return True

    text = str(value).strip()

    return text in ["", "nan", "None", "/", "--", "-"]



def _exclude_word_and_branch_summary_branches(df: pd.DataFrame) -> pd.DataFrame:
    """
    排除不参与附录分公司统计和 Word 统计/图表的特殊分公司。
    注意：只在汇总统计处调用，不影响名单、交流会成绩明细、未出勤名单等原始明细输出。
    """
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return df

    data = df.copy()

    if "分公司" not in data.columns:
        return data

    branch_series = data["分公司"].astype(str).str.strip()
    return data[~branch_series.isin(EXCLUDED_WORD_AND_BRANCH_SUMMARY_BRANCHES)].copy()


def _extract_exam_date_from_submit_times(submit_times, file_path: str = ""):
    """
    从 Excel 的“交卷时间”列获取考试当天日期。

    规则：读取所有人的“交卷时间”，取最早的有效交卷时间，其日期即为该科考试日期。
    例如最早交卷时间为 2026-05-28 10:20:43，则考试日期为 2026-05-28。
    """
    submit_dt = pd.to_datetime(submit_times, errors="coerce")
    valid_dt = submit_dt.dropna()

    if valid_dt.empty:
        raise ValueError(
            f"无法从‘交卷时间’列获取考试日期：{file_path}。请确认该列至少有一个有效时间。"
        )

    return valid_dt.min().date()


def _extract_meeting_month_label_from_filename(file_path: str) -> str:
    """
    从成绩文件名中提取交流会月份。
    示例：考试考生情况-2026年5月中国区技术交流会考试--化免20260608114705.xlsx -> 5月份
    """
    filename = os.path.basename(str(file_path))

    patterns = [
        r"\d{4}年\s*(\d{1,2})月\s*中国区技术交流会考试",
        r"\d{4}年\s*(\d{1,2})月[^\d]*中国区技术交流会",
        r"(\d{1,2})月\s*中国区技术交流会",
    ]

    for pattern in patterns:
        match = re.search(pattern, filename)
        if match:
            return f"{int(match.group(1))}月份"

    return ""


def _get_meeting_month_label_from_exam_files(exam_files: Dict[str, str]) -> str:
    """从六科成绩文件中提取交流会月份，取第一个能识别到的月份。"""
    if not exam_files:
        return ""

    for file_path in exam_files.values():
        label = _extract_meeting_month_label_from_filename(file_path)
        if label:
            return label

    return ""


def _normalize_channel_unknown_branch(df: pd.DataFrame) -> pd.DataFrame:
    """
    渠道员工无法解析分公司时，统一写入“未知分公司”。
    这样后续名单、成绩表、出勤率统计和异常名单都能保持一致。
    """
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return df

    data = df.copy()

    if "工号" not in data.columns:
        data["工号"] = ""

    if "人员类型" not in data.columns:
        data["人员类型"] = data["工号"].apply(
            lambda x: "渠道" if str(x).strip().upper().startswith("S") else "内部"
        )

    if "分公司" not in data.columns:
        data["分公司"] = ""

    channel_mask = data["人员类型"].astype(str).str.strip().eq("渠道")
    blank_branch_mask = data["分公司"].apply(_is_blank_value)
    data.loc[channel_mask & blank_branch_mask, "分公司"] = UNKNOWN_BRANCH_NAME

    return data


def _get_unknown_channel_branch_records(df: pd.DataFrame) -> pd.DataFrame:
    """
    提取渠道员工分公司无法解析的异常名单。
    """
    output_cols = ["工号", "姓名", "分公司", "渠道商", "部门", "人员类型"]

    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return pd.DataFrame(columns=output_cols)

    data = df.copy()

    for col in output_cols:
        if col not in data.columns:
            data[col] = ""

    if "人员类型" not in data.columns:
        data["人员类型"] = data["工号"].apply(
            lambda x: "渠道" if str(x).strip().upper().startswith("S") else "内部"
        )

    mask = (
        data["人员类型"].astype(str).str.strip().eq("渠道")
        & data["分公司"].astype(str).str.strip().eq(UNKNOWN_BRANCH_NAME)
    )

    result = data.loc[mask, output_cols].copy()
    result["异常原因"] = "渠道员工部门路径未解析到分公司，已在相关表格中填写未知分公司"
    result = result.drop_duplicates(subset=["工号"], keep="first")
    return result


def _normalize_subject_list(value) -> List[str]:
    """
    将 present_subjects / attended_subjects 统一转成科目列表。
    present_subjects 表示“这个人出现在该科原始成绩表中”，不受考试状态、交卷时间影响。
    """
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]

    if isinstance(value, tuple) or isinstance(value, set):
        return [str(x).strip() for x in value if str(x).strip()]

    if pd.isna(value):
        return []

    text = str(value).strip()
    if text in ["", "[]", "nan", "None", "/"]:
        return []

    # 兼容被 Excel/日志转成字符串的列表，例如 "['化免', 'TLA']"
    text = text.strip("[]")
    parts = re.split(r"[,，;；]", text)
    return [p.strip().strip("\'\"") for p in parts if p.strip().strip("\'\"")]


def _get_row_present_subjects(row) -> List[str]:
    """
    获取某人的“实际出现在原始成绩表中的科目”。
    优先使用 present_subjects；旧数据没有该字段时，才兼容 attended_subjects。
    """
    present_subjects = _normalize_subject_list(row.get("present_subjects", []))
    if present_subjects:
        return present_subjects
    return _normalize_subject_list(row.get("attended_subjects", []))


def _calculate_average_from_present_subjects(row, exam_types: List[str]) -> float:
    """
    平均成绩口径：只计算“原始成绩表中出现过”的科目。
    出现过且成绩为 0，必须参与平均；没有出现的科目“/”不参与平均。
    """
    present_subjects = _get_row_present_subjects(row)
    scores = []

    for subject in exam_types:
        if subject not in present_subjects:
            continue
        score = pd.to_numeric(row.get(subject, 0), errors="coerce")
        scores.append(0.0 if pd.isna(score) else float(score))

    return round(sum(scores) / len(scores), 1) if scores else 0.0


def _drop_rows_without_present_subjects(df: pd.DataFrame, exam_types: List[str]) -> pd.DataFrame:
    """
    剔除六科都没有出现在原始成绩表中的异常人员。
    正常情况下名单来自六科成绩表，不应出现六科全“/”人员；这里作为兜底。
    """
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return df

    data = df.copy()

    if "present_subjects" in data.columns:
        mask = data.apply(lambda row: len(_get_row_present_subjects(row)) > 0, axis=1)
        return data.loc[mask].copy()

    # 没有 present_subjects 时，按展示值兜底判断：至少有一个科目不是“/”才保留。
    for subject in exam_types:
        if subject not in data.columns:
            data[subject] = "/"

    mask = data[exam_types].apply(
        lambda row: any(str(v).strip() not in ["", "/", "nan", "None"] for v in row),
        axis=1
    )
    return data.loc[mask].copy()


def _apply_absent_subject_slash_for_output(df: pd.DataFrame, exam_types: List[str]) -> pd.DataFrame:
    """
    仅用于输出展示：没有出现在某科原始成绩表中的人员，该科显示“/”。
    只要这个人在某科原始成绩表中出现过，即使考试状态为未完成、未提交，成绩为 0，也显示 0。
    """
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return df

    data = df.copy()

    for subject in exam_types:
        if subject not in data.columns:
            data[subject] = "/"
            continue

        data[subject] = data[subject].astype(object)

        def display_score(row):
            present_subjects = _get_row_present_subjects(row)
            if subject in present_subjects:
                value = row.get(subject, 0)
                if pd.isna(value):
                    return 0
                return value
            return "/"

        data[subject] = data.apply(display_score, axis=1)

    data = _drop_rows_without_present_subjects(data, exam_types)
    return data




def _apply_attendance_by_score_logic(df: pd.DataFrame, exam_types: Optional[List[str]] = None) -> pd.DataFrame:
    """
    统一出勤口径：
    不出勤 = 六科成绩均为 0 或没有成绩；
    实际出勤 = 总人数 - 不出勤人数。

    说明：这里不再以“交卷时间/attended_subjects”判断出勤，
    而是以最终进入汇总的六科成绩是否至少有一科大于 0 判断。
    """
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return df

    if exam_types is None:
        exam_types = ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]

    data = df.copy()

    for subject in exam_types:
        if subject not in data.columns:
            data[subject] = 0

    score_df = data[exam_types].apply(lambda col: pd.to_numeric(col, errors="coerce")).fillna(0)

    # 至少一个科目成绩 > 0，才算实际出勤。
    data["attended"] = score_df.gt(0).any(axis=1)
    data["是否不出勤"] = ~data["attended"]

    return data


def _build_branch_summary_df(data: pd.DataFrame, staff_type: str, exam_types: List[str]) -> pd.DataFrame:
    """
    按人员类型生成分公司出勤率及平均成绩统计。

    重要口径：
    1. 应出勤人数 = 该分公司纳入统计的总人数；
    2. 实际出勤人数 = 至少有一个科目成绩 > 0 的人数；
    3. 平均成绩 = 该分公司“实际出勤人员”的个人平均成绩再求平均。
       不出勤人员（所有成绩均为0或没有成绩）不参与分公司平均成绩计算。

    Word 内部/渠道考核成绩图不再重新计算分公司成绩，而是直接读取本函数写入
    “中国区技术交流会出勤和考核情况-IVD线.xlsx”中
    “分公司出勤率以及平均成绩”sheet 的平均成绩列。
    """
    columns = ["分公司", "应出勤人数", "实际出勤人数", "出勤率", "平均成绩"]

    if data is None or data.empty:
        return pd.DataFrame(columns=columns)

    temp = data[data["人员类型"] == staff_type].copy()

    if temp.empty:
        return pd.DataFrame(columns=columns)

    # 新出勤口径：所有成绩均为 0 或没有成绩的人算不出勤；否则算实际出勤。
    temp = _apply_attendance_by_score_logic(temp, exam_types)

    # 确保每个人都有“平均成绩”：按原始成绩表中出现过的科目计算，0分参与平均，“/”不参与平均。
    temp["平均成绩"] = temp.apply(
        lambda row: _calculate_average_from_present_subjects(row, exam_types),
        axis=1
    )

    rows = []

    for branch, group in temp.groupby("分公司", dropna=False, sort=False):
        branch_name = UNKNOWN_BRANCH_NAME if _is_blank_value(branch) else str(branch).strip()

        should_attend = int(len(group))
        actual_attend = int(group["attended"].sum()) if "attended" in group.columns else 0
        attendance_rate = actual_attend / should_attend if should_attend else 0

        # 分公司平均成绩只统计实际出勤人员。
        attended_group = group[group["attended"] == True].copy() if "attended" in group.columns else group.iloc[0:0].copy()
        avg_values = pd.to_numeric(attended_group["平均成绩"], errors="coerce").dropna()
        avg_score = round(float(avg_values.mean()), 1) if len(avg_values) > 0 else 0.0

        rows.append({
            "分公司": branch_name,
            "应出勤人数": should_attend,
            "实际出勤人数": actual_attend,
            "出勤率": attendance_rate,
            "平均成绩": avg_score
        })

    return pd.DataFrame(rows, columns=columns).sort_values("分公司").reset_index(drop=True)



def _read_branch_summary_sheet_for_word(excel_path: str) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    读取“中国区技术交流会出勤和考核情况-IVD线.xlsx”中的
    “分公司出勤率以及平均成绩”sheet。

    Word中的四个图表直接使用这个sheet的数据：
    - 内部/渠道出勤率图：分公司、应出勤人数、实际出勤人数、出勤率；
    - 内部/渠道考核成绩图：分公司、实际出勤人数、平均成绩。

    这样可以保证Word图表与Excel附录中的分公司统计完全一致，
    不再在Word绘图处重新按人员明细计算第二套成绩逻辑。
    """
    columns = ["分公司", "应出勤人数", "实际出勤人数", "出勤率", "平均成绩"]
    empty = pd.DataFrame(columns=columns)

    if not excel_path or not os.path.exists(excel_path):
        print(f"⚠️ 未找到分公司统计Excel，Word图表将无法直接读取附录数据：{excel_path}")
        return empty.copy(), empty.copy()

    try:
        wb = load_workbook(excel_path, data_only=True, read_only=True)
    except Exception as e:
        print(f"⚠️ 打开分公司统计Excel失败，Word图表将无法直接读取附录数据：{e}")
        return empty.copy(), empty.copy()

    try:
        sheet_name = "分公司出勤率以及平均成绩"
        if sheet_name not in wb.sheetnames:
            print(f"⚠️ Excel中不存在sheet：{sheet_name}，Word图表将无法直接读取附录数据")
            return empty.copy(), empty.copy()

        ws = wb[sheet_name]

        def read_side_table(start_col: int) -> pd.DataFrame:
            rows = []
            row_idx = 3  # 第1行标题，第2行表头，第3行开始数据

            while row_idx <= ws.max_row:
                values = [ws.cell(row=row_idx, column=start_col + i).value for i in range(5)]

                # 侧表遇到全空行即可停止。内部和渠道分别读取，因此互不影响。
                if all(v is None or str(v).strip() == "" for v in values):
                    break

                branch = values[0]
                if branch is None or str(branch).strip() in ["", "nan", "None", "/"]:
                    row_idx += 1
                    continue

                rows.append({
                    "分公司": str(branch).strip(),
                    "应出勤人数": values[1],
                    "实际出勤人数": values[2],
                    "出勤率": values[3],
                    "平均成绩": values[4],
                })
                row_idx += 1

            df = pd.DataFrame(rows, columns=columns)

            if df.empty:
                return empty.copy()

            for col in ["应出勤人数", "实际出勤人数", "平均成绩"]:
                df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

            def parse_rate(value):
                if pd.isna(value):
                    return 0.0
                if isinstance(value, str):
                    text = value.strip().replace("%", "")
                    number = pd.to_numeric(text, errors="coerce")
                    if pd.isna(number):
                        return 0.0
                    return float(number) / 100 if "%" in value else float(number)
                number = pd.to_numeric(value, errors="coerce")
                return 0.0 if pd.isna(number) else float(number)

            df["出勤率"] = df["出勤率"].apply(parse_rate)
            df["应出勤人数"] = df["应出勤人数"].round(0).astype(int)
            df["实际出勤人数"] = df["实际出勤人数"].round(0).astype(int)
            df["平均成绩"] = df["平均成绩"].round(1)

            return df

        internal_df = read_side_table(1)
        channel_df = read_side_table(8)

        print(
            f"✅ Word图表已读取附录分公司统计sheet：内部{len(internal_df)}行，渠道{len(channel_df)}行"
        )
        return internal_df, channel_df

    finally:
        try:
            wb.close()
        except Exception:
            pass


def enrich_person_info_from_exam_dfs(
    merged_data: pd.DataFrame,
    all_exam_dfs: List[pd.DataFrame]
) -> pd.DataFrame:
    """
    从六科成绩表中汇总人员基础信息，并回填到 merged_data。

    解决：
    1. 内部未出勤人员分公司为空；
    2. 渠道未出勤人员分公司/渠道商为空；
    3. 避免 KeyError: '_parsed_branch'。

    回填字段：
    姓名、部门、分公司、渠道商、人员类型
    """

    if merged_data is None:
        return merged_data

    if not isinstance(merged_data, pd.DataFrame):
        return merged_data

    if merged_data.empty:
        return merged_data

    info_parts = []

    for exam_df in all_exam_dfs:
        if exam_df is None:
            continue

        if not isinstance(exam_df, pd.DataFrame):
            continue

        if exam_df.empty:
            continue

        temp = exam_df.copy()

        if "工号" not in temp.columns and "exam_id" in temp.columns:
            temp["工号"] = temp["exam_id"]

        if "工号" not in temp.columns:
            continue

        for col in ["姓名", "部门", "分公司", "渠道商", "人员类型"]:
            if col not in temp.columns:
                temp[col] = ""

        temp["工号"] = temp["工号"].astype(str).str.strip()

        # 关键：统一从部门字段补齐分公司、渠道商
        temp = _fill_branch_channel_from_department(temp)

        info_parts.append(
            temp[["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]]
        )

    if not info_parts:
        return merged_data

    info_df = pd.concat(info_parts, ignore_index=True)

    info_df["工号"] = info_df["工号"].astype(str).str.strip()
    info_df = info_df[info_df["工号"] != ""].copy()

    # 优先保留信息更完整的一条
    info_df["_info_score"] = (
        info_df["姓名"].astype(str).str.len()
        + info_df["部门"].astype(str).str.len()
        + info_df["分公司"].astype(str).str.len() * 2
        + info_df["渠道商"].astype(str).str.len() * 3
    )

    info_df = info_df.sort_values(
        by=["工号", "_info_score"],
        ascending=[True, False]
    )

    info_df = info_df.drop_duplicates(subset=["工号"], keep="first")
    info_df = info_df.drop(columns=["_info_score"])

    data = merged_data.copy()

    if "工号" not in data.columns:
        return data

    data["工号"] = data["工号"].astype(str).str.strip()

    for col in ["姓名", "部门", "分公司", "渠道商", "人员类型"]:
        if col not in data.columns:
            data[col] = ""

    data = data.merge(
        info_df,
        on="工号",
        how="left",
        suffixes=("", "_from_exam")
    )

    for col in ["姓名", "部门", "分公司", "渠道商", "人员类型"]:
        from_col = f"{col}_from_exam"

        if from_col in data.columns:
            data[col] = data.apply(
                lambda row: row[from_col]
                if _is_blank_value(row.get(col)) and not _is_blank_value(row.get(from_col))
                else row.get(col),
                axis=1
            )

            data = data.drop(columns=[from_col])

    # 再兜底一次：如果 merged_data 自身有部门，则继续解析分公司、渠道商
    data = _fill_branch_channel_from_department(data)

    if "人员类型" not in data.columns:
        data["人员类型"] = data["工号"].apply(
            lambda x: "渠道" if str(x).startswith("S") else "内部"
        )
    else:
        data["人员类型"] = data.apply(
            lambda row: "渠道"
            if str(row.get("工号")).startswith("S")
            else row.get("人员类型"),
            axis=1
        )

    # 内部员工渠道商置空
    data.loc[data["人员类型"] == "内部", "渠道商"] = ""

    return data


def _calculate_adjusted_score(raw_score, submit_time, dispatch_date, exam_status=None):
    """
    成绩修正规则：
    1. 派发当天考试：最终成绩 = 原始成绩
    2. 派发日期之后考试：
       - 原始成绩 >= 80，最终成绩 = 80
       - 原始成绩 < 80，最终成绩 = 原始成绩
    3. 未开始/未提交/无交卷时间：最终成绩 = 0
    """
    raw_score = _to_numeric_score(raw_score)

    status_text = "" if pd.isna(exam_status) else str(exam_status).strip()

    if status_text in ["未开始", "未提交"]:
        return 0.0

    submit_dt = pd.to_datetime(submit_time, errors="coerce")

    if pd.isna(submit_dt):
        return 0.0

    submit_date = submit_dt.date()

    if submit_date == dispatch_date:
        return raw_score

    if submit_date > dispatch_date:
        return 80.0 if raw_score >= 80 else raw_score

    return raw_score

def load_exam_scores(exam_type: str, file_path: str) -> pd.DataFrame:
    """
    加载单科考试成绩并标准化。

    新版规则：
    1. 六科均通过“账号”获取唯一 ID；
    2. 从“部门”解析分公司、渠道商；
    3. 从“成绩”列获取原始成绩；
    4. 从Excel交卷时间列获取考试日期；
    5. 派发当天考试保留原始成绩；
       派发日期之后考试，成绩 >= 80 则最终成绩为 80，否则保留原成绩；
    6. 返回字段兼容旧版 merge_attendance_with_scores：
       exam_type、exam_id、score。
    """

    print(f"正在读取 {exam_type} 成绩文件：{file_path}")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"{exam_type} 成绩文件不存在：{file_path}")

    header_row = _find_header_row(
        file_path,
        required_cols=["账号", "部门", "成绩"]
    )

    df = pd.read_excel(
        file_path,
        header=header_row,
        dtype=str,
        engine="openpyxl"
    )

    df.columns = [str(col).strip() for col in df.columns]
    df = df.dropna(how="all").copy()

    required_columns = ["账号", "姓名", "部门", "成绩", "交卷时间"]

    missing_columns = [col for col in required_columns if col not in df.columns]

    if missing_columns:
        raise ValueError(
            f"{exam_type} 成绩文件缺少必要列：{missing_columns}，当前识别列：{list(df.columns)}"
        )

    dispatch_date = _extract_exam_date_from_submit_times(df["交卷时间"], file_path)

    df["工号"] = df["账号"].apply(_normalize_account)
    df = df[df["工号"] != ""].copy()

    df[["分公司", "渠道商"]] = df["部门"].apply(
        lambda x: pd.Series(_parse_department_info(x))
    )

    df["人员类型"] = df["工号"].apply(
        lambda x: "渠道" if str(x).startswith("S") else "内部"
    )

    df["原始成绩"] = df["成绩"].apply(_to_numeric_score)

    if "考试状态" in df.columns:
        df[exam_type] = df.apply(
            lambda row: _calculate_adjusted_score(
                raw_score=row.get("成绩"),
                submit_time=row.get("交卷时间"),
                dispatch_date=dispatch_date,
                exam_status=row.get("考试状态")
            ),
            axis=1
        )
    else:
        df[exam_type] = df.apply(
            lambda row: _calculate_adjusted_score(
                raw_score=row.get("成绩"),
                submit_time=row.get("交卷时间"),
                dispatch_date=dispatch_date,
                exam_status=None
            ),
            axis=1
        )

    submit_dt = pd.to_datetime(df["交卷时间"], errors="coerce")

    if "考试状态" in df.columns:
        status = df["考试状态"].fillna("").astype(str).str.strip()
        attended_mask = submit_dt.notna() & (~status.isin(["未开始", "未提交"]))
    else:
        attended_mask = submit_dt.notna()

    df["attended_subjects"] = attended_mask.apply(
        lambda x: [exam_type] if x else []
    )

    # 关键字段：只要该人员出现在该科原始成绩表中，就记录为 present。
    # 后续“/”展示和平均成绩计算必须用 present_subjects，不能用考试状态/交卷时间判断。
    df["present_subjects"] = df["工号"].apply(lambda _: [exam_type])

    keep_columns = [
        "工号",
        "姓名",
        "部门",
        "分公司",
        "渠道商",
        "人员类型",
        "原始成绩",
        exam_type,
        "attended_subjects",
        "present_subjects",
        "交卷时间"
    ]

    optional_columns = ["考试状态", "通过状态", "账号状态", "岗位"]

    for col in optional_columns:
        if col in df.columns and col not in keep_columns:
            keep_columns.append(col)

    result = df[keep_columns].copy()

    result["exam_type"] = exam_type
    result["exam_id"] = result["工号"]
    result["score"] = result[exam_type]

    result["_attended_flag"] = result["attended_subjects"].apply(
        lambda x: 1 if x else 0
    )

    result = result.sort_values(
        by=["工号", "_attended_flag", exam_type],
        ascending=[True, False, False]
    )

    result = result.drop_duplicates(subset=["工号"], keep="first")
    result = result.drop(columns=["_attended_flag"])

    result["exam_type"] = exam_type
    result["exam_id"] = result["工号"]
    result["score"] = result[exam_type]

    print(
        f"{exam_type} 成绩读取完成：共 {len(result)} 人，"
        f"考试日期：{dispatch_date}，"
        f"出现在成绩表人数：{result['present_subjects'].apply(lambda x: bool(x)).sum()}，"
        f"有效交卷人数：{result['attended_subjects'].apply(lambda x: bool(x)).sum()}"
    )

    return result


# 【其余函数保持不变】

def _fill_branch_channel_from_department(df: pd.DataFrame) -> pd.DataFrame:
    """
    从“部门”字段补齐“分公司”和“渠道商”。

    用途：
    1. build_roster_from_exam_dfs()
    2. enrich_person_info_from_exam_dfs()
    3. generate_absentee_lists()
    4. create_appendix_workbook() 的兜底处理
    """

    if df is None:
        return df

    if not isinstance(df, pd.DataFrame):
        return df

    if df.empty:
        return df

    data = df.copy()

    if "部门" not in data.columns:
        return data

    if "分公司" not in data.columns:
        data["分公司"] = ""

    if "渠道商" not in data.columns:
        data["渠道商"] = ""

    parsed = data["部门"].apply(lambda x: pd.Series(_parse_department_info(x)))
    parsed.columns = ["_parsed_branch", "_parsed_channel"]

    branch_mask = data["分公司"].apply(_is_blank_value) & parsed["_parsed_branch"].apply(
        lambda x: not _is_blank_value(x)
    )

    channel_mask = data["渠道商"].apply(_is_blank_value) & parsed["_parsed_channel"].apply(
        lambda x: not _is_blank_value(x)
    )

    data.loc[branch_mask, "分公司"] = parsed.loc[branch_mask, "_parsed_branch"]
    data.loc[channel_mask, "渠道商"] = parsed.loc[channel_mask, "_parsed_channel"]

    data = _normalize_channel_unknown_branch(data)

    return data

def build_roster_from_exam_dfs(all_exam_dfs: List[pd.DataFrame]) -> pd.DataFrame:
    """
    从六个成绩 DataFrame 中汇总所有人员，去重生成名单。

    输出字段：
    工号、姓名、部门、分公司、渠道商、人员类型

    说明：
    1. 不再读取名单.xlsx；
    2. 名单来自六科成绩表所有账号去重；
    3. 从“部门”字段补齐分公司、渠道商；
    4. 不再使用 row["_parsed_branch"]，避免 KeyError。
    """

    roster_parts = []

    for exam_df in all_exam_dfs:
        if exam_df is None:
            continue

        if not isinstance(exam_df, pd.DataFrame):
            print(f"警告：发现非 DataFrame 类型成绩数据，已跳过：{type(exam_df)}")
            continue

        if exam_df.empty:
            continue

        temp = exam_df.copy()

        if "工号" not in temp.columns and "exam_id" in temp.columns:
            temp["工号"] = temp["exam_id"]

        if "工号" not in temp.columns:
            print(f"警告：成绩数据缺少工号列，已跳过。当前列：{list(temp.columns)}")
            continue

        for col in ["姓名", "部门", "分公司", "渠道商"]:
            if col not in temp.columns:
                temp[col] = ""

        # 关键：统一从部门字段补齐分公司、渠道商
        temp = _fill_branch_channel_from_department(temp)

        if "人员类型" not in temp.columns:
            temp["人员类型"] = temp["工号"].apply(
                lambda x: "渠道" if str(x).startswith("S") else "内部"
            )

        temp["工号"] = temp["工号"].astype(str).str.strip()

        roster_parts.append(
            temp[["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]]
        )

    if not roster_parts:
        return pd.DataFrame(
            columns=["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]
        )

    roster_df = pd.concat(roster_parts, ignore_index=True)

    roster_df["工号"] = roster_df["工号"].astype(str).str.strip()
    roster_df = roster_df[roster_df["工号"] != ""].copy()

    # 信息完整度评分：同一个人在多个科目出现时，保留信息最完整的记录
    roster_df["_info_score"] = (
        roster_df["姓名"].astype(str).str.len()
        + roster_df["部门"].astype(str).str.len()
        + roster_df["分公司"].astype(str).str.len() * 2
        + roster_df["渠道商"].astype(str).str.len() * 3
    )

    roster_df = roster_df.sort_values(
        by=["工号", "_info_score"],
        ascending=[True, False]
    )

    roster_df = roster_df.drop_duplicates(subset=["工号"], keep="first")
    roster_df = roster_df.drop(columns=["_info_score"])

    # 内部员工不应显示渠道商
    roster_df.loc[roster_df["人员类型"] == "内部", "渠道商"] = ""

    return roster_df


def merge_attendance_with_scores(
    attendance_roster: pd.DataFrame,
    all_exam_dfs: List[pd.DataFrame]
) -> pd.DataFrame:
    """
    关联名单与六科成绩。

    注意：
    现在 attendance_roster 不再来自名单.xlsx，
    而是由六科成绩中的所有人员去重汇总生成。
    """

    exam_types = ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]

    base = pd.DataFrame({
        "工号": attendance_roster["工号"].astype(str),
        "姓名": attendance_roster["姓名"] if "姓名" in attendance_roster.columns else "",
        "部门": attendance_roster["部门"] if "部门" in attendance_roster.columns else "",
        "分公司": attendance_roster["分公司"] if "分公司" in attendance_roster.columns else "",
        "渠道商": attendance_roster["渠道商"] if "渠道商" in attendance_roster.columns else "",
        "人员类型": attendance_roster["人员类型"] if "人员类型" in attendance_roster.columns else [
            "渠道" if str(x).startswith("S") else "内部" for x in attendance_roster["工号"]
        ]
    })

    base["attended_subjects"] = [[] for _ in range(len(base))]
    base["present_subjects"] = [[] for _ in range(len(base))]

    for exam_type in exam_types:
        base[exam_type] = 0.0

    for exam_df in all_exam_dfs:
        if exam_df is None or exam_df.empty:
            print("警告：发现空的科目成绩数据，已跳过。")
            continue

        exam_df = exam_df.copy()

        if "exam_type" in exam_df.columns:
            et = exam_df["exam_type"].iloc[0]
        else:
            matched = [s for s in exam_types if s in exam_df.columns]
            if not matched:
                raise ValueError(
                    f"无法识别科目，当前列为：{list(exam_df.columns)}"
                )
            et = matched[0]
            exam_df["exam_type"] = et

        if "exam_id" not in exam_df.columns:
            if "工号" in exam_df.columns:
                exam_df["exam_id"] = exam_df["工号"]
            else:
                raise ValueError(
                    f"成绩数据缺少 exam_id，且没有工号列，当前列为：{list(exam_df.columns)}"
                )

        if "score" not in exam_df.columns:
            if et in exam_df.columns:
                exam_df["score"] = exam_df[et]
            else:
                raise ValueError(
                    f"成绩数据缺少 score，且没有科目成绩列 {et}，当前列为：{list(exam_df.columns)}"
                )

        merge_cols = ["exam_id", "score"]

        if "attended_subjects" in exam_df.columns:
            merge_cols.append("attended_subjects")

        if "present_subjects" in exam_df.columns:
            merge_cols.append("present_subjects")

        temp = exam_df[merge_cols].copy()
        rename_map = {
            "exam_id": "工号",
            "score": et,
            "attended_subjects": f"{et}_attended_subjects",
            "present_subjects": f"{et}_present_subjects"
        }
        temp = temp.rename(columns=rename_map)

        temp["工号"] = temp["工号"].astype(str)

        base = base.merge(
            temp,
            on="工号",
            how="left",
            suffixes=("", "_new")
        )

        if f"{et}_new" in base.columns:
            base[et] = base[f"{et}_new"].fillna(base[et])
            base = base.drop(columns=[f"{et}_new"])

        base[et] = base[et].fillna(0).astype(float)

        def merge_subjects(old_list, new_list):
            result = []
            result.extend(_normalize_subject_list(old_list))
            result.extend(_normalize_subject_list(new_list))
            return list(dict.fromkeys(result))

        att_col = f"{et}_attended_subjects"

        if att_col in base.columns:
            base["attended_subjects"] = base.apply(
                lambda row: merge_subjects(row["attended_subjects"], row[att_col]),
                axis=1
            )

            base = base.drop(columns=[att_col])

        present_col = f"{et}_present_subjects"

        if present_col in base.columns:
            base["present_subjects"] = base.apply(
                lambda row: merge_subjects(row["present_subjects"], row[present_col]),
                axis=1
            )

            base = base.drop(columns=[present_col])

    score_cols = exam_types

    # 新出勤口径：所有成绩均为 0 或没有成绩的人算不出勤；否则算实际出勤。
    base = _apply_attendance_by_score_logic(base, score_cols)


    base["平均成绩"] = base.apply(
        lambda row: _calculate_average_from_present_subjects(row, score_cols),
        axis=1
    )
    base["最高成绩"] = base[score_cols].max(axis=1)
    base["是否通过"] = base[score_cols].ge(80).any(axis=1)

    return base


def calculate_overall_summary(merged_data: pd.DataFrame) -> Dict:
    """生成整体概况指标。Word 总结口径会排除指定特殊分公司。"""
    exam_types = ['凝血', '化免', '临检', 'TLA', '微生物', '尿液']
    data = _exclude_word_and_branch_summary_branches(merged_data)
    if data is None or data.empty:
        total_attendees = 0
        actual_attended = 0
        passed = 0
    else:
        data = _apply_attendance_by_score_logic(data, exam_types)
        total_attendees = len(data)
        absent_people = int((~data['attended']).sum())
        actual_attended = total_attendees - absent_people
        passed = (data[exam_types] >= 80).any(axis=1).sum()
    pass_rate = (passed / actual_attended * 100) if actual_attended > 0 else 0.0

    status_dist = {
        "已完成": actual_attended,
        "未开始": total_attendees - actual_attended,
        "归档": 0
    }

    return {
        "总参与人数": total_attendees,
        "实际参考人数": actual_attended,
        "整体通过率": round(pass_rate, 1),
        "状态分布": status_dist
    }


def generate_absentee_lists(merged_data: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    """
    生成未出勤名单。

    输出：
    1. 内部未出勤：工号、姓名、分公司
    2. 渠道未出勤：工号、姓名、分公司、渠道商

    同时返回中文 key 和英文 key，兼容旧代码：
    - 内部未出勤 / internal_absentees
    - 渠道未出勤 / channel_absentees
    """

    exam_types = ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]

    if merged_data is None or merged_data.empty:
        empty_internal = pd.DataFrame(columns=["工号", "姓名", "分公司"])
        empty_channel = pd.DataFrame(columns=["工号", "姓名", "分公司", "渠道商"])

        return {
            "内部未出勤": empty_internal,
            "渠道未出勤": empty_channel,
            "internal_absentees": empty_internal,
            "channel_absentees": empty_channel
        }

    data = merged_data.copy()

    for col in exam_types:
        if col not in data.columns:
            data[col] = 0

    for col in ["工号", "姓名", "部门", "分公司", "渠道商"]:
        if col not in data.columns:
            data[col] = ""

    if "人员类型" not in data.columns:
        data["人员类型"] = data["工号"].apply(
            lambda x: "渠道" if str(x).startswith("S") else "内部"
        )

    # 关键：统一补齐分公司和渠道商
    data = _fill_branch_channel_from_department(data)

    # 内部员工渠道商置空
    data.loc[data["人员类型"] == "内部", "渠道商"] = ""

    # 新出勤口径：所有成绩均为 0 或没有成绩的人算不出勤。
    data = _apply_attendance_by_score_logic(data, exam_types)

    absentees = data[~data["attended"]].copy()

    internal_cols = ["工号", "姓名", "分公司"]
    channel_cols = ["工号", "姓名", "分公司", "渠道商"]

    for col in internal_cols:
        if col not in absentees.columns:
            absentees[col] = ""

    for col in channel_cols:
        if col not in absentees.columns:
            absentees[col] = ""

    internal_absentees = absentees[
        absentees["人员类型"] == "内部"
    ][internal_cols].copy()

    channel_absentees = absentees[
        absentees["人员类型"] == "渠道"
    ][channel_cols].copy()

    return {
        "内部未出勤": internal_absentees,
        "渠道未出勤": channel_absentees,
        "internal_absentees": internal_absentees,
        "channel_absentees": channel_absentees
    }

def read_exam_excel_for_output(file_path: str) -> pd.DataFrame:
    """
    用于写入“交流会统计结果.xlsx”的各科原始成绩 sheet。

    自动跳过前面的考试标题行，避免出现：
    2026年05月凝血技术交流会考试
    2026年05月凝血技术交流会考试.1
    ...

    新增：在原始“成绩”列后一列插入“最终成绩”。
    规则：
    1. Excel中最早有效交卷时间对应日期为考试日期；
    2. 派发当天交卷，最终成绩 = 原始成绩；
    3. 派发日期之后交卷，原始成绩 >= 80 时最终成绩 = 80，原始成绩 < 80 时最终成绩 = 原始成绩；
    4. 未开始/未提交/无交卷时间时，最终成绩 = 0。
    """

    header_row = _find_header_row(
        file_path,
        required_cols=["账号", "部门", "成绩"]
    )

    df = pd.read_excel(
        file_path,
        header=header_row,
        dtype=str,
        engine="openpyxl"
    )

    df.columns = [str(c).strip() for c in df.columns]

    df = df.dropna(axis=1, how="all")
    df = df.dropna(axis=0, how="all")

    df = df.loc[:, ~df.columns.astype(str).str.startswith("Unnamed")]

    if "成绩" not in df.columns:
        raise ValueError(f"成绩文件缺少‘成绩’列，无法计算最终成绩：{file_path}")

    if "交卷时间" not in df.columns:
        raise ValueError(f"成绩文件缺少‘交卷时间’列，无法根据交卷时间计算最终成绩：{file_path}")

    dispatch_date = _extract_exam_date_from_submit_times(df["交卷时间"], file_path)

    # 避免重复运行时出现旧的“最终成绩”列。
    if "最终成绩" in df.columns:
        df = df.drop(columns=["最终成绩"])

    status_col = "考试状态" if "考试状态" in df.columns else None

    df["最终成绩"] = df.apply(
        lambda row: _calculate_adjusted_score(
            raw_score=row.get("成绩"),
            submit_time=row.get("交卷时间"),
            dispatch_date=dispatch_date,
            exam_status=row.get(status_col) if status_col else None
        ),
        axis=1
    )

    # 将“最终成绩”移动到“成绩”后一列。
    columns = list(df.columns)
    columns.remove("最终成绩")
    score_index = columns.index("成绩")
    columns.insert(score_index + 1, "最终成绩")
    df = df[columns]

    return df


def create_main_statistics_file(
    roster_file: str = "",
    exam_files: dict = None,
    output_path: str = "交流会统计结果.xlsx",
    merged_data: pd.DataFrame = None
):
    """
    完全使用openpyxl创建Excel文件，避免pandas与openpyxl混合使用的问题
    """
    wb = Workbook()

    # 删除默认sheet
    default_sheet = wb.active
    wb.remove(default_sheet)

    if merged_data is not None:
        merged_data = _normalize_channel_unknown_branch(_fill_branch_channel_from_department(merged_data))
        merged_data = _apply_attendance_by_score_logic(merged_data, ["凝血", "化免", "临检", "TLA", "微生物", "尿液"])

    # 1. 复制名单数据到'名单'工作表
    # =====================================================
    # 新规则：现在不再强制读取名单.xlsx
    # 如果 roster_file 为空，则从 merged_data 中生成名单
    # =====================================================
    if roster_file and str(roster_file).strip() and os.path.exists(roster_file):
        roster_df = pd.read_excel(roster_file, dtype={"工号": str}, engine="openpyxl")
    else:
        roster_cols = ["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]

        temp = merged_data.copy()

        for col in roster_cols:
            if col not in temp.columns:
                temp[col] = ""

        roster_df = temp[roster_cols].copy()
        roster_df["工号"] = roster_df["工号"].astype(str).str.strip()
        roster_df = roster_df[roster_df["工号"] != ""].copy()
        roster_df = roster_df.drop_duplicates(subset=["工号"], keep="first")
    ws_roster = wb.create_sheet(title='名单')

    # 写入表头
    headers = roster_df.columns.tolist()
    for col_num, header in enumerate(headers, 1):
        cell = ws_roster.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 写入数据
    for row_num, row_data in enumerate(roster_df.values, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            ws_roster.cell(row=row_num, column=col_num, value=cell_value)

    # 2. 复制各科成绩数据到对应的工作表
    for subject, file_path in exam_files.items():
        try:
            exam_df = read_exam_excel_for_output(file_path)
            ws_subject = wb.create_sheet(title=subject)

            # 写入表头
            headers = exam_df.columns.tolist()
            for col_num, header in enumerate(headers, 1):
                cell = ws_subject.cell(row=1, column=col_num, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")

            # 写入数据
            for row_num, row_data in enumerate(exam_df.values, 2):
                for col_num, cell_value in enumerate(row_data, 1):
                    ws_subject.cell(row=row_num, column=col_num, value=cell_value)
        except Exception as e:
            print(f"加载{subject}科目数据失败: {str(e)}")

    # 3. 创建"内部成绩"sheet页
    print("创建内部成绩sheet...")
    internal_df = merged_data[merged_data['人员类型'] == '内部'].copy()
    exam_types = ['化免', '临检', 'TLA', '凝血', '微生物', '尿液']
    internal_df = _drop_rows_without_present_subjects(internal_df, exam_types)

    # 按照要求的顺序排列列。注意：present_subjects 只用于判断该科是否出现在原始成绩表中，写入前会删除。
    internal_stats = pd.DataFrame({
        '工号': internal_df['工号'],
        '姓名': internal_df['姓名'],
        '分公司': internal_df['分公司'],
        '化免': internal_df['化免'],
        '临检': internal_df['临检'],
        'TLA': internal_df['TLA'],
        '凝血': internal_df['凝血'],
        '微生物': internal_df['微生物'],
        '尿液': internal_df['尿液'],
        'present_subjects': internal_df['present_subjects'] if 'present_subjects' in internal_df.columns else internal_df.get('attended_subjects', [[] for _ in range(len(internal_df))])
    })

    # 只要该人员出现在某科原始成绩表中，该科就显示实际成绩；成绩为0也显示0。
    # 只有未出现在该科原始成绩表中，该科才显示“/”。
    internal_stats = _apply_absent_subject_slash_for_output(internal_stats, exam_types)

    # 平均成绩：按原始成绩表中出现过的科目计算，0分参与平均，“/”不参与平均。
    internal_stats['平均成绩'] = internal_stats.apply(
        lambda row: _calculate_average_from_present_subjects(row, exam_types),
        axis=1
    )
    internal_stats = internal_stats.drop(columns=['present_subjects'], errors='ignore')
    internal_stats = internal_stats.sort_values('工号').reset_index(drop=True)

    # 内部分公司平均成绩统计
    internal_dept_avg = []
    for dept, group in internal_df.groupby('分公司'):
        if len(group) == 0:
            continue
        dept_scores = []
        for _, row in group.iterrows():
            present_subjects = _get_row_present_subjects(row)
            if not present_subjects:
                continue
            # 计算该员工出现在原始成绩表中的科目平均分，0分参与平均。
            dept_scores.append(_calculate_average_from_present_subjects(row, exam_types))
        if dept_scores:
            internal_dept_avg.append({
                '分公司': dept,
                '平均成绩': round(sum(dept_scores) / len(dept_scores), 1),
                '人数': len(group)
            })
    dept_avg_df = pd.DataFrame(internal_dept_avg).sort_values(by='平均成绩', ascending=False).reset_index(
        drop=True)

    # 创建内部成绩工作表
    ws_internal = wb.create_sheet(title='内部成绩')

    # 写入内部成绩主体数据表头
    internal_headers = list(internal_stats.columns)
    for col_num, header in enumerate(internal_headers, 1):
        cell = ws_internal.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 写入内部成绩主体数据
    for row_num, row_data in enumerate(internal_stats.values, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            ws_internal.cell(row=row_num, column=col_num, value=cell_value)
            # 科目列居中对齐
            if internal_headers[col_num-1] in exam_types:
                ws_internal.cell(row=row_num, column=col_num).alignment = Alignment(horizontal="center")

    # 添加空行分隔
    separator_row = len(internal_stats) + 3

    # 写入分公司统计标题
    title_cell = ws_internal.cell(row=separator_row, column=1, value="内部员工分公司平均成绩统计")
    title_cell.font = Font(bold=True, size=14)

    # 写入分公司统计表头
    dept_headers = list(dept_avg_df.columns)
    dept_header_row = separator_row + 2
    for col_num, header in enumerate(dept_headers, 1):
        cell = ws_internal.cell(row=dept_header_row, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 写入分公司统计数据
    for row_num, row_data in enumerate(dept_avg_df.values, dept_header_row + 1):
        for col_num, cell_value in enumerate(row_data, 1):
            ws_internal.cell(row=row_num, column=col_num, value=cell_value)

    print(f"内部成绩sheet创建完成，共{len(internal_stats)}条记录，{len(dept_avg_df)}个分公司")

    # 4. 创建"渠道成绩"sheet页
    print("创建渠道成绩sheet...")
    channel_df = merged_data[merged_data['人员类型'] == '渠道'].copy()
    channel_df = _drop_rows_without_present_subjects(channel_df, exam_types)

    # 按照要求的顺序排列列。注意：present_subjects 只用于判断该科是否出现在原始成绩表中，写入前会删除。
    channel_stats = pd.DataFrame({
        '工号': channel_df['工号'],
        '姓名': channel_df['姓名'],
        '分公司': channel_df['分公司'],
        '化免': channel_df['化免'],
        '临检': channel_df['临检'],
        'TLA': channel_df['TLA'],
        '凝血': channel_df['凝血'],
        '微生物': channel_df['微生物'],
        '尿液': channel_df['尿液'],
        'present_subjects': channel_df['present_subjects'] if 'present_subjects' in channel_df.columns else channel_df.get('attended_subjects', [[] for _ in range(len(channel_df))])
    })

    # 只要该人员出现在某科原始成绩表中，该科就显示实际成绩；成绩为0也显示0。
    # 只有未出现在该科原始成绩表中，该科才显示“/”。
    channel_stats = _apply_absent_subject_slash_for_output(channel_stats, exam_types)

    # 平均成绩：按原始成绩表中出现过的科目计算，0分参与平均，“/”不参与平均。
    channel_stats['平均成绩'] = channel_stats.apply(
        lambda row: _calculate_average_from_present_subjects(row, exam_types),
        axis=1
    )
    channel_stats = channel_stats.drop(columns=['present_subjects'], errors='ignore')
    channel_stats = channel_stats.sort_values('工号').reset_index(drop=True)

    # 渠道分公司平均成绩统计
    channel_dept_avg = []
    for dept, group in channel_df.groupby('分公司'):
        if len(group) == 0:
            continue
        dept_scores = []
        for _, row in group.iterrows():
            present_subjects = _get_row_present_subjects(row)
            if not present_subjects:
                continue
            # 计算该员工出现在原始成绩表中的科目平均分，0分参与平均。
            dept_scores.append(_calculate_average_from_present_subjects(row, exam_types))
        if dept_scores:
            channel_dept_avg.append({
                '分公司': dept,
                '平均成绩': round(sum(dept_scores) / len(dept_scores), 1),
                '人数': len(group)
            })
    channel_dept_avg_df = pd.DataFrame(channel_dept_avg).sort_values(by='平均成绩',
                                                                     ascending=False).reset_index(
        drop=True)

    # 创建渠道成绩工作表
    ws_channel = wb.create_sheet(title='渠道成绩')

    # 写入渠道成绩主体数据表头
    channel_headers = list(channel_stats.columns)
    for col_num, header in enumerate(channel_headers, 1):
        cell = ws_channel.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 写入渠道成绩主体数据
    for row_num, row_data in enumerate(channel_stats.values, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            ws_channel.cell(row=row_num, column=col_num, value=cell_value)
            # 科目列居中对齐
            if channel_headers[col_num-1] in exam_types:
                ws_channel.cell(row=row_num, column=col_num).alignment = Alignment(horizontal="center")

    # 添加空行分隔
    channel_separator_row = len(channel_stats) + 3

    # 写入渠道分公司统计标题
    channel_title_cell = ws_channel.cell(row=channel_separator_row, column=1, value="渠道员工分公司平均成绩统计")
    channel_title_cell.font = Font(bold=True, size=14)

    # 写入渠道分公司统计表头
    channel_dept_headers = list(channel_dept_avg_df.columns)
    channel_dept_header_row = channel_separator_row + 2
    for col_num, header in enumerate(channel_dept_headers, 1):
        cell = ws_channel.cell(row=channel_dept_header_row, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # 写入渠道分公司统计数据
    for row_num, row_data in enumerate(channel_dept_avg_df.values, channel_dept_header_row + 1):
        for col_num, cell_value in enumerate(row_data, 1):
            ws_channel.cell(row=row_num, column=col_num, value=cell_value)

    print(f"渠道成绩sheet创建完成，共{len(channel_stats)}条记录，{len(channel_dept_avg_df)}个分公司")

    # ===================== 新增：内部出勤率sheet =====================
    print("创建内部出勤率sheet...")
    # 计算内部人员分公司出勤率
    internal_attendance_data = []
    internal_data = merged_data[merged_data['人员类型'] == '内部']
    for dept, group in internal_data.groupby('分公司'):
        should_attend = len(group)  # 应出勤人数（名单中的总人数）
        actual_attend = group['attended'].sum()  # 实际出勤人数（任意一科成绩>0）
        attendance_rate = (actual_attend / should_attend * 100) if should_attend > 0 else 0.0
        internal_attendance_data.append({
            '分公司': dept,
            '实际出勤': actual_attend,
            '应出勤': should_attend,
            '出勤率': f"{round(attendance_rate, 1)}%"
        })
    # 转为DataFrame并按分公司排序
    internal_attendance_df = pd.DataFrame(internal_attendance_data).sort_values('分公司').reset_index(drop=True)

    # 创建内部出勤率工作表
    ws_internal_att = wb.create_sheet(title='内部出勤率')
    # 写入表头
    internal_att_headers = ['分公司', '实际出勤', '应出勤', '出勤率']
    for col_num, header in enumerate(internal_att_headers, 1):
        cell = ws_internal_att.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    # 写入数据
    for row_num, row_data in enumerate(internal_attendance_df.values, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            cell = ws_internal_att.cell(row=row_num, column=col_num, value=cell_value)
            # 出勤率列居中对齐
            if col_num == 4:
                cell.alignment = Alignment(horizontal="center", vertical="center")
    print(f"内部出勤率sheet创建完成，共{len(internal_attendance_df)}个分公司")

    # ===================== 新增：渠道出勤率sheet =====================
    print("创建渠道出勤率sheet...")
    # 计算渠道人员分公司出勤率
    channel_attendance_data = []
    channel_data = merged_data[merged_data['人员类型'] == '渠道']
    for dept, group in channel_data.groupby('分公司'):
        should_attend = len(group)  # 应出勤人数（名单中的总人数）
        actual_attend = group['attended'].sum()  # 实际出勤人数（任意一科成绩>0）
        attendance_rate = (actual_attend / should_attend * 100) if should_attend > 0 else 0.0
        channel_attendance_data.append({
            '分公司': dept,
            '实际出勤': actual_attend,
            '应出勤': should_attend,
            '出勤率': f"{round(attendance_rate, 1)}%"
        })
    # 转为DataFrame并按分公司排序
    channel_attendance_df = pd.DataFrame(channel_attendance_data).sort_values('分公司').reset_index(drop=True)

    # 创建渠道出勤率工作表
    ws_channel_att = wb.create_sheet(title='渠道出勤率')
    # 写入表头
    channel_att_headers = ['分公司', '实际出勤', '应出勤', '出勤率']
    for col_num, header in enumerate(channel_att_headers, 1):
        cell = ws_channel_att.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    # 写入数据
    for row_num, row_data in enumerate(channel_attendance_df.values, 2):
        for col_num, cell_value in enumerate(row_data, 1):
            cell = ws_channel_att.cell(row=row_num, column=col_num, value=cell_value)
            # 出勤率列居中对齐
            if col_num == 4:
                cell.alignment = Alignment(horizontal="center", vertical="center")
    print(f"渠道出勤率sheet创建完成，共{len(channel_attendance_df)}个分公司")


    # 保存文件
    wb.save(output_path)
    wb.close()  # 确保文件被正确关闭
    print(f"文件已成功保存至: {output_path}")

"""生成附录Excel文件（适配新需求）"""


def create_appendix_workbook(
    merged_data: pd.DataFrame,
    absentee_lists: Dict[str, pd.DataFrame],
    roster_df: pd.DataFrame,
    output_path: str
):
    """
    生成附录 Excel 文件。

    Sheet：
    1. 名单汇总
    2. 交流会成绩
    3. 分公司出勤率以及平均成绩
    4. 内部未出勤
    5. 渠道未出勤

    修复：
    1. 内部未出勤补齐分公司；
    2. 渠道未出勤补齐分公司、渠道商；
    3. 不再使用 row["_parsed_branch"]；
    4. 兼容中文/英文 absentee_lists key。
    """

    exam_types = ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]

    if merged_data is None:
        merged_data = pd.DataFrame()

    if roster_df is None:
        roster_df = pd.DataFrame()

    merged_data = merged_data.copy()
    roster_df = roster_df.copy()

    # 兜底补齐基础列
    for df in [merged_data, roster_df]:
        for col in ["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]:
            if col not in df.columns:
                df[col] = ""

    merged_data = _normalize_channel_unknown_branch(_fill_branch_channel_from_department(merged_data))
    merged_data = _apply_attendance_by_score_logic(merged_data, exam_types)
    roster_df = _normalize_channel_unknown_branch(_fill_branch_channel_from_department(roster_df))

    if "人员类型" not in merged_data.columns:
        merged_data["人员类型"] = merged_data["工号"].apply(
            lambda x: "渠道" if str(x).startswith("S") else "内部"
        )

    if "人员类型" not in roster_df.columns:
        roster_df["人员类型"] = roster_df["工号"].apply(
            lambda x: "渠道" if str(x).startswith("S") else "内部"
        )

    merged_data.loc[merged_data["人员类型"] == "内部", "渠道商"] = ""
    roster_df.loc[roster_df["人员类型"] == "内部", "渠道商"] = ""

    wb = Workbook()

    default_sheet = wb.active
    wb.remove(default_sheet)

    header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    header_font = Font(bold=True)
    center_align = Alignment(horizontal="center", vertical="center")

    def write_df_to_sheet(ws, df: pd.DataFrame):
        """
        通用 DataFrame 写入 Sheet。
        """
        if df is None:
            df = pd.DataFrame()

        df = df.copy()

        headers = list(df.columns)

        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align

        for row_num, row_data in enumerate(df.values, 2):
            for col_num, cell_value in enumerate(row_data, 1):
                ws.cell(row=row_num, column=col_num, value=cell_value)

        # 简单自适应列宽
        for col_idx, col_name in enumerate(headers, 1):
            max_len = len(str(col_name))
            for value in df[col_name].astype(str).head(500):
                max_len = max(max_len, len(value))

            width = min(max(max_len + 2, 12), 45)
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width

    # =====================================================
    # 1. 名单汇总
    # =====================================================
    ws_roster = wb.create_sheet(title="名单汇总")

    roster_cols = ["工号", "姓名", "分公司", "渠道商", "人员类型", "部门"]

    for col in roster_cols:
        if col not in roster_df.columns:
            roster_df[col] = ""

    roster_out = roster_df[roster_cols].copy()
    roster_out["工号"] = roster_out["工号"].astype(str).str.strip()
    roster_out = roster_out[roster_out["工号"] != ""].copy()
    roster_out = roster_out.drop_duplicates(subset=["工号"], keep="first")

    write_df_to_sheet(ws_roster, roster_out)

    # =====================================================
    # 2. 交流会成绩
    # =====================================================
    ws_score = wb.create_sheet(title="交流会成绩")

    score_cols = ["工号", "姓名", "分公司", "渠道商", "人员类型"] + exam_types + ["平均成绩"]

    for col in score_cols:
        if col not in merged_data.columns:
            if col in exam_types:
                merged_data[col] = 0
            else:
                merged_data[col] = ""

    helper_cols = []
    if "present_subjects" in merged_data.columns:
        helper_cols.append("present_subjects")
    if "attended_subjects" in merged_data.columns:
        helper_cols.append("attended_subjects")

    score_out = merged_data[score_cols + helper_cols].copy()
    score_out = _apply_absent_subject_slash_for_output(score_out, exam_types)

    # 重新计算展示用平均成绩：0分参与平均，“/”不参与平均。
    score_out["平均成绩"] = score_out.apply(
        lambda row: _calculate_average_from_present_subjects(row, exam_types),
        axis=1
    )

    score_out = score_out.drop(columns=["present_subjects", "attended_subjects"], errors="ignore")

    write_df_to_sheet(ws_score, score_out)

    # =====================================================
    # 3. 分公司出勤率以及平均成绩
    # =====================================================
    ws_branch = wb.create_sheet(title="分公司出勤率以及平均成绩")

    data_for_branch = _exclude_word_and_branch_summary_branches(
        _normalize_channel_unknown_branch(merged_data.copy())
    )

    data_for_branch = _apply_attendance_by_score_logic(data_for_branch, exam_types)

    if "平均成绩" not in data_for_branch.columns:
        data_for_branch["平均成绩"] = data_for_branch[exam_types].replace(0, np.nan).mean(axis=1).fillna(0)

    internal_branch_out = _build_branch_summary_df(data_for_branch, "内部", exam_types)
    channel_branch_out = _build_branch_summary_df(data_for_branch, "渠道", exam_types)

    def write_side_by_side_branch_summary():
        internal_title = ws_branch.cell(row=1, column=1, value="内部员工分公司出勤率以及平均成绩")
        internal_title.font = Font(bold=True, size=12)
        channel_title = ws_branch.cell(row=1, column=8, value="渠道员工分公司出勤率以及平均成绩")
        channel_title.font = Font(bold=True, size=12)

        def write_table(start_col: int, df: pd.DataFrame):
            headers = ["分公司", "应出勤人数", "实际出勤人数", "出勤率", "平均成绩"]
            for idx, header in enumerate(headers, start_col):
                cell = ws_branch.cell(row=2, column=idx, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align

            if df is None or df.empty:
                return

            for row_num, row_data in enumerate(df[headers].values, 3):
                for offset, cell_value in enumerate(row_data):
                    cell = ws_branch.cell(row=row_num, column=start_col + offset, value=cell_value)
                    if headers[offset] == "出勤率":
                        cell.number_format = "0.0%"
                        cell.alignment = center_align

            for offset, header in enumerate(headers):
                col_idx = start_col + offset
                max_len = len(header)
                for value in df[header].astype(str).head(500):
                    max_len = max(max_len, len(value))
                ws_branch.column_dimensions[ws_branch.cell(row=2, column=col_idx).column_letter].width = min(max(max_len + 2, 12), 45)

        write_table(1, internal_branch_out)
        write_table(8, channel_branch_out)

    write_side_by_side_branch_summary()

    # =====================================================
    # 4. 内部未出勤
    # =====================================================
    ws_internal_abs = wb.create_sheet(title="内部未出勤")

    internal_abs_df = absentee_lists.get("内部未出勤")
    if internal_abs_df is None:
        internal_abs_df = absentee_lists.get("internal_absentees", pd.DataFrame())

    internal_abs_df = internal_abs_df.copy()

    # 从 roster/merged 信息回填
    person_info = pd.concat(
        [
            roster_df[["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]],
            merged_data[["工号", "姓名", "部门", "分公司", "渠道商", "人员类型"]]
        ],
        ignore_index=True
    )

    person_info = _fill_branch_channel_from_department(person_info)
    person_info["工号"] = person_info["工号"].astype(str).str.strip()

    person_info["_info_score"] = (
        person_info["姓名"].astype(str).str.len()
        + person_info["部门"].astype(str).str.len()
        + person_info["分公司"].astype(str).str.len() * 2
        + person_info["渠道商"].astype(str).str.len() * 3
    )

    person_info = person_info.sort_values(
        by=["工号", "_info_score"],
        ascending=[True, False]
    ).drop_duplicates(subset=["工号"], keep="first")

    person_info = person_info.drop(columns=["_info_score"])

    def fill_absentee_info(abs_df: pd.DataFrame, output_cols: List[str]) -> pd.DataFrame:
        if abs_df is None:
            abs_df = pd.DataFrame(columns=output_cols)

        abs_df = abs_df.copy()

        for col in output_cols:
            if col not in abs_df.columns:
                abs_df[col] = ""

        if abs_df.empty:
            return pd.DataFrame(columns=output_cols)

        abs_df["工号"] = abs_df["工号"].astype(str).str.strip()

        abs_df = abs_df.merge(
            person_info,
            on="工号",
            how="left",
            suffixes=("", "_补充")
        )

        for col in ["姓名", "分公司", "渠道商"]:
            supplement_col = f"{col}_补充"

            if col not in abs_df.columns:
                abs_df[col] = ""

            if supplement_col in abs_df.columns:
                abs_df[col] = abs_df.apply(
                    lambda row: row[supplement_col]
                    if _is_blank_value(row.get(col)) and not _is_blank_value(row.get(supplement_col))
                    else row.get(col),
                    axis=1
                )

        # 如果有部门补充，再从部门解析一次
        if "部门_补充" in abs_df.columns:
            temp_parse = abs_df.rename(columns={"部门_补充": "部门"}).copy()
            temp_parse = _fill_branch_channel_from_department(temp_parse)

            if "分公司" in temp_parse.columns:
                abs_df["分公司"] = abs_df.apply(
                    lambda row: temp_parse.loc[row.name, "分公司"]
                    if _is_blank_value(row.get("分公司")) and not _is_blank_value(temp_parse.loc[row.name, "分公司"])
                    else row.get("分公司"),
                    axis=1
                )

            if "渠道商" in output_cols and "渠道商" in temp_parse.columns:
                abs_df["渠道商"] = abs_df.apply(
                    lambda row: temp_parse.loc[row.name, "渠道商"]
                    if _is_blank_value(row.get("渠道商")) and not _is_blank_value(temp_parse.loc[row.name, "渠道商"])
                    else row.get("渠道商"),
                    axis=1
                )

        for col in output_cols:
            if col not in abs_df.columns:
                abs_df[col] = ""

        return abs_df[output_cols].copy()

    internal_abs_cols = ["工号", "姓名", "分公司"]
    internal_abs_data = fill_absentee_info(internal_abs_df, internal_abs_cols)

    write_df_to_sheet(ws_internal_abs, internal_abs_data)

    # =====================================================
    # 5. 渠道未出勤
    # =====================================================
    ws_channel_abs = wb.create_sheet(title="渠道未出勤")

    channel_abs_df = absentee_lists.get("渠道未出勤")
    if channel_abs_df is None:
        channel_abs_df = absentee_lists.get("channel_absentees", pd.DataFrame())

    channel_abs_cols = ["工号", "姓名", "分公司", "渠道商"]
    channel_abs_data = fill_absentee_info(channel_abs_df, channel_abs_cols)

    write_df_to_sheet(ws_channel_abs, channel_abs_data)

    # =====================================================
    # 6. 渠道分公司异常名单
    # =====================================================
    ws_channel_unknown = wb.create_sheet(title="渠道分公司异常名单")
    unknown_channel_df = _get_unknown_channel_branch_records(merged_data)
    write_df_to_sheet(ws_channel_unknown, unknown_channel_df)

    # 保存
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    wb.save(output_path)
    wb.close()

    print(f"附录文件已成功保存至: {output_path}")






# 解决中文乱码
plt.rcParams["font.family"] = ["SimHei", "Microsoft YaHei"]
plt.rcParams["axes.unicode_minus"] = False


# ===================== 全局配置：中文渲染 + 图表样式 =====================
plt.rcParams["font.family"] = ["SimHei", "Microsoft YaHei"]
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["figure.dpi"] = 300  # 高清图表
# 标准配色（严格匹配需求）
COLOR_BLUE = "#4472C4"  # 应出勤/成绩柱形
COLOR_ORANGE = "#ED7D31"  # 实际出勤/成绩折线
COLOR_GRAY = "#7F7F7F"  # 出勤率折线


# ===================== 新增1：读取出勤考核Excel数据 =====================



# ===================== 新增1：读取出勤考核Excel数据 (修复版) =====================


def read_attendance_score_excel(file_path: str) -> Dict[str, Optional[pd.DataFrame]]:
    """
    读取技术交流会出勤成绩Excel文件，增加了对Excel错误值和特殊格式的容错处理。
    """
    result = {
        "内部成绩": None,
        "渠道成绩": None,
        "内部出勤率": None,
        "渠道出勤率": None,
        "交流会成绩": None,
        "分公司出勤率及平均成绩": None,
        "内部未出勤": None,
        "渠道未出勤": None
    }

    try:
        # 1. 获取Sheet名称
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        sheet_names = wb.sheetnames
        wb.close()

        # 2. 逐个读取Sheet并清洗
        # --- 内部成绩 ---
        if "内部成绩" in sheet_names:
            # 读取时不指定特定列的类型，保持原始object类型，便于后续处理
            df_internal_score = pd.read_excel(file_path, sheet_name="内部成绩", na_filter=True)
            result["内部成绩"] = _clean_score_data_fixed(df_internal_score, is_internal=True)

        # --- 渠道成绩 ---
        if "渠道成绩" in sheet_names:
            df_channel_score = pd.read_excel(file_path, sheet_name="渠道成绩", na_filter=True)
            result["渠道成绩"] = _clean_score_data_fixed(df_channel_score, is_internal=False)

        # --- 内部出勤率 ---
        if "内部出勤率" in sheet_names:
            df_internal_att = pd.read_excel(file_path, sheet_name="内部出勤率", na_filter=True)
            # 处理出勤率列中的错误值和百分比
            if "出勤率" in df_internal_att.columns:
                # 将错误值替换为 NaN，然后转为数值
                df_internal_att["出勤率"] = (
                    df_internal_att["出勤率"]
                    .replace(['#N/A', '#DIV/0!', '#VALUE!', '#REF!', '#NAME?', '/'], np.nan)
                    .astype(str)
                    .str.replace('%', '', regex=False) # 移除百分号
                )
                # 强制转换为数值，无法转换的变为 NaN，然后填充为 0
                df_internal_att["出勤率"] = pd.to_numeric(df_internal_att["出勤率"], errors='coerce').fillna(0.0)
            result["内部出勤率"] = df_internal_att

        # --- 渠道出勤率 ---
        if "渠道出勤率" in sheet_names:
            df_channel_att = pd.read_excel(file_path, sheet_name="渠道出勤率", na_filter=True)
            if "出勤率" in df_channel_att.columns:
                 df_channel_att["出勤率"] = (
                    df_channel_att["出勤率"]
                    .replace(['#N/A', '#DIV/0!', '#VALUE!', '#REF!', '#NAME?', '/'], np.nan)
                    .astype(str)
                    .str.replace('%', '', regex=False)
                )
            else:
                df_channel_att["出勤率"] = pd.to_numeric(df_channel_att["出勤率"], errors='coerce').fillna(0.0)
            result["渠道出勤率"] = df_channel_att

        # --- 交流会成绩 ---
        if "交流会成绩" in sheet_names:
            df_exchange_score = pd.read_excel(file_path, sheet_name="交流会成绩", na_filter=True)
            result["交流会成绩"] = _clean_score_data_fixed(df_exchange_score, is_internal=None)

        # --- 分公司出勤率以及平均成绩 ---
        if "分公司出勤率以及平均成绩" in sheet_names:
            df_dept_summary = pd.read_excel(file_path, sheet_name="分公司出勤率以及平均成绩", na_filter=True)
            # 此Sheet可能包含合并单元格或复杂结构，这里仅做基本读取，不深入处理
            result["分公司出勤率及平均成绩"] = df_dept_summary

        # --- 未出勤名单 ---
        for sheet_name in ["内部未出勤", "渠道未出勤"]:
            if sheet_name in sheet_names:
                df_absent = pd.read_excel(file_path, sheet_name=sheet_name, na_filter=True)
                key = "内部未出勤" if sheet_name == "内部未出勤" else "渠道未出勤"
                # 提取工号、姓名、分公司三列，并去除完全为空的行
                if all(col in df_absent.columns for col in ["工号", "姓名", "分公司"]):
                     result[key] = df_absent[["工号", "姓名", "分公司"]].dropna(how='all')
                else:
                    # 如果列名不匹配，尝试寻找最接近的列或记录警告
                    print(f"警告: {sheet_name} Sheet中缺少必要的列 ['工号', '姓名', '分公司']，跳过读取。")

    except Exception as e:
        print(f"读取出勤成绩Excel失败: {str(e)}，跳过相关统计")
        # 即使出错也返回空结构，防止后续代码崩溃
        return result

    return result


def _clean_score_data_fixed(df: pd.DataFrame, is_internal: Optional[bool] = None) -> pd.DataFrame:
    """
    清洗成绩数据的修复版。
    更加健壮地处理各种错误值和特殊字符。
    """
    if df.empty:
        return df

    # 1. 定义核心列
    core_cols = ["工号", "姓名", "分公司", "化免", "临检", "TLA", "凝血", "微生物", "尿液", "平均成绩"]
    # 只选择存在的列
    valid_cols = [col for col in core_cols if col in df.columns]
    df = df[valid_cols].copy()

    # 2. 处理考试科目成绩
    exam_subjects = ["化免", "临检", "TLA", "凝血", "微生物", "尿液"]
    for subject in exam_subjects:
        if subject in df.columns:
            # 将常见的错误值和特殊符号替换为 NaN
            df[subject] = df[subject].replace(['/', '#N/A', '#DIV/0!', '#VALUE!', '#REF!', '#NAME?'], np.nan)
            # 强制转换为数值，无法转换的变为 NaN
            df[subject] = pd.to_numeric(df[subject], errors='coerce')
            # 将 NaN 填充为 0.0
            df[subject] = df[subject].fillna(0.0)

    # 3. 处理平均成绩
    if "平均成绩" in df.columns:
        # 同样处理错误值和百分号
        df["平均成绩"] = df["平均成绩"].replace(['/', '#N/A', '#DIV/0!', '#VALUE!', '#REF!', '#NAME?'], np.nan)
        df["平均成绩"] = pd.to_numeric(
            df["平均成绩"].astype(str).str.replace('%', '', regex=False), errors='coerce'
        ).fillna(0.0).round(1)

    # 4. 清洗工号和分公司
    if "工号" in df.columns:
        # 将工号转换为字符串，去除首尾空格，处理未知值
        df["工号"] = df["工号"].astype(str).str.strip().str.upper().replace({'UNKNOWN': '0', 'nan': '0', 'None': '0'})

    if "分公司" in df.columns:
        # 清洗分公司名称
        df["分公司"] = df["分公司"].astype(str).str.strip().replace({'nan': '未知', 'None': '未知'})

    # 5. 排序
    if "工号" in df.columns:
        df = df.sort_values("工号").reset_index(drop=True)

    return df

# ... (其余代码保持不变)

def _clean_branch_name(x):
    """清洗分公司名称，去掉空值、总计行等"""
    if pd.isna(x):
        return None
    x = str(x).strip()
    if x in ["", "/", "nan", "None", "总计", "合计"]:
        return None
    return x


def _plot_combo_bar_line(
    df: pd.DataFrame,
    branch_col: str,
    bar1_col: str,
    bar2_col: str,
    line_col: str,
    save_path: str,
    title: str,
    bar1_label: str,
    bar2_label: str,
    line_label: str,
    y1_label: str = "",
    y2_label: str = "",
    line_is_percent: bool = True,
    figsize=(13.5, 4.6)
):
    """
    生成类似Excel组合图的图片：
    - 蓝色柱：应出勤 / 人数 / 成绩基准
    - 橙色柱：实际出勤 / 平均成绩
    - 灰色折线：出勤率 / 达标率 / 其他百分比
    - 右轴显示百分比
    """

    if df is None or df.empty:
        print(f"⚠️ {title} 数据为空，跳过绘图")
        return

    plot_df = df.copy()

    # 清洗分公司
    plot_df[branch_col] = plot_df[branch_col].apply(_clean_branch_name)
    plot_df = plot_df.dropna(subset=[branch_col])

    # 数值列转换
    for col in [bar1_col, bar2_col, line_col]:
        plot_df[col] = pd.to_numeric(plot_df[col], errors="coerce")

    plot_df = plot_df.dropna(subset=[bar1_col, bar2_col, line_col])

    if plot_df.empty:
        print(f"⚠️ {title} 清洗后无有效数据，跳过绘图")
        return

    # 出勤率如果是 0~1，转成 0~100
    if line_is_percent and plot_df[line_col].max() <= 1.5:
        plot_df[line_col] = plot_df[line_col] * 100

    # 按原顺序展示，不强制排序，更接近你截图里的区域排列
    branches = plot_df[branch_col].tolist()
    x = np.arange(len(branches))
    width = 0.34

    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax1 = plt.subplots(figsize=figsize, dpi=300)

    # 柱形图
    bar1 = ax1.bar(
        x - width / 2,
        plot_df[bar1_col],
        width,
        label=bar1_label,
        color="#5B9BD5"
    )

    bar2 = ax1.bar(
        x + width / 2,
        plot_df[bar2_col],
        width,
        label=bar2_label,
        color="#ED7D31"
    )

    # 左轴样式
    ax1.set_ylabel(y1_label, fontsize=9)
    ax1.set_xticks(x)
    ax1.set_xticklabels(branches, rotation=45, ha="right", fontsize=8)
    ax1.tick_params(axis="y", labelsize=8, colors="#666666")
    ax1.tick_params(axis="x", labelsize=8, colors="#666666")
    ax1.grid(axis="y", color="#D9D9D9", linewidth=0.8)
    ax1.set_axisbelow(True)

    # 去掉边框，保留Excel干净风格
    for spine in ["top", "right", "left", "bottom"]:
        ax1.spines[spine].set_visible(False)

    # 右轴折线
    ax2 = ax1.twinx()
    line = ax2.plot(
        x,
        plot_df[line_col],
        label=line_label,
        color="#A5A5A5",
        marker="o",
        markersize=4,
        linewidth=2.2
    )

    ax2.set_ylabel(y2_label, fontsize=9)
    ax2.tick_params(axis="y", labelsize=8, colors="#666666")

    if line_is_percent:
        max_rate = max(120, int(np.ceil(plot_df[line_col].max() / 20) * 20))
        ax2.set_ylim(0, max_rate)
        ax2.set_yticks(np.arange(0, max_rate + 1, 20))
        ax2.set_yticklabels([f"{int(v)}%" for v in np.arange(0, max_rate + 1, 20)])
    else:
        ax2.set_ylim(0, max(100, plot_df[line_col].max() * 1.15))

    for spine in ["top", "right", "left", "bottom"]:
        ax2.spines[spine].set_visible(False)

    # 折线顶部标签
    for i, val in enumerate(plot_df[line_col]):
        if pd.notna(val):
            label = f"{val:.0f}%" if line_is_percent else f"{val:.1f}"
            ax2.text(
                x[i],
                val + (2 if line_is_percent else 1),
                label,
                ha="center",
                va="bottom",
                fontsize=8,
                color="#595959"
            )

    # 标题
    ax1.set_title(title, fontsize=14, color="#595959", pad=15)

    # 图例放底部居中
    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()

    ax1.legend(
        handles1 + handles2,
        labels1 + labels2,
        loc="upper center",
        bbox_to_anchor=(0.5, -0.22),
        ncol=3,
        frameon=False,
        fontsize=8
    )

    # 横向留白
    ax1.margins(x=0.01)

    plt.tight_layout()
    plt.savefig(save_path, bbox_inches="tight", dpi=300)
    plt.close()


# ===================== 重写：生成总结Word文档（核心函数） =====================

warnings.filterwarnings('ignore')




def generate_summary_docx(
    merged_data: pd.DataFrame,
    absentee_lists: Dict[str, pd.DataFrame],
    overall_summary: Dict,
    roster_df: pd.DataFrame,
    template_path: str = "中国区技术交流会总结邮件内容.docx",
    output_path: str = "中国区技术交流会总结邮件.docx",
    meeting_month_label: str = "",
    branch_summary_excel_path: str = ""
):
    """
    基于模板生成技术交流会总结docx文档（完美保留原格式版）
    1. 读取"中国区技术交流会总结邮件内容.docx"模板
    2. 保留原格式替换应出勤人数、实际出勤人数、总出勤率
    3. 在指定位置精准插入4个统计图表
    """
    # 1. 读取模板文档（完整保留原有内容、格式和签名）
    try:
        doc = Document(template_path)
    except FileNotFoundError:
        raise FileNotFoundError(f"模板文件不存在，请检查文件路径：{template_path}")

    # Word 总结与图表统计口径：排除指定特殊分公司，并按新出勤口径重算。
    # 这里同时统一“未知分公司”的写法，确保 Word 图表与附录“分公司出勤率以及平均成绩”sheet 口径一致。
    exam_types_for_attendance = ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]
    merged_data = _normalize_channel_unknown_branch(merged_data)
    merged_data = _exclude_word_and_branch_summary_branches(merged_data)
    merged_data = _apply_attendance_by_score_logic(merged_data, exam_types_for_attendance)
    overall_summary = calculate_overall_summary(merged_data)
    title_prefix = str(meeting_month_label or "").strip()

    # Word图表数据源：直接读取附录Excel“分公司出勤率以及平均成绩”sheet，
    # 不再在Word绘图处重新计算第二套分公司平均成绩逻辑。
    branch_internal_summary_df, branch_channel_summary_df = _read_branch_summary_sheet_for_word(
        branch_summary_excel_path
    )

    # 兜底：如果外部调用没有传入附录Excel路径，才回退到同一个函数生成的分公司汇总数据。
    # 正常GUI/主流程会传入“中国区技术交流会出勤和考核情况-IVD线.xlsx”，
    # 因此Word成绩图会直接使用附录sheet中的“平均成绩”列。
    if branch_internal_summary_df.empty:
        branch_internal_summary_df = _build_branch_summary_df(merged_data, "内部", exam_types_for_attendance)
    if branch_channel_summary_df.empty:
        branch_channel_summary_df = _build_branch_summary_df(merged_data, "渠道", exam_types_for_attendance)

    # 2. 计算核心统计数据
    total_should_attend = overall_summary["总参与人数"]
    total_actual_attend = overall_summary["实际参考人数"]
    total_attendance_rate = round((total_actual_attend / total_should_attend) * 100,
                                  1) if total_should_attend > 0 else 0.0
    new_attendance_text = f"本次交流会应出勤{total_should_attend}人，实际出勤{total_actual_attend}人（含回看），总出勤率{total_attendance_rate}%。"

    # ===================== 核心修复：保留原格式的文本替换 =====================
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # 匹配包含出勤统计的段落（精确匹配整行，避免误替换）
        if "本次交流会应出勤" in text and "实际出勤" in text and "总出勤率" in text:
            # 步骤1：保存原段落的所有格式属性
            if paragraph.runs:
                # 取第一个run的格式作为基准（模板中整行格式统一）
                original_run = paragraph.runs[0]
                font_name = original_run.font.name
                font_size = original_run.font.size
                font_bold = original_run.font.bold
                font_color = original_run.font.color.rgb if original_run.font.color else None
                alignment = paragraph.alignment
            else:
                # 极端情况：段落无run，使用模板默认格式（微软雅黑小四）
                font_name = "微软雅黑"
                font_size = Pt(12)  # 小四对应12磅
                font_bold = False
                font_color = None
                alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 步骤2：清空段落所有内容（保留段落本身的格式）
            paragraph.clear()

            # 步骤3：按原格式写入新文本
            new_run = paragraph.add_run(new_attendance_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color
            # 设置中文字体（解决微软雅黑不生效的问题）
            new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            paragraph.alignment = alignment

            # 找到目标段落后立即退出循环，避免重复替换
            break

    # 3. 定义图表生成辅助函数（内存中生成，不产生临时文件）
    def generate_chart_to_bytesio(plot_func):
        """生成图表并返回BytesIO二进制流"""
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
        plt.rcParams['axes.unicode_minus'] = False

        fig, ax = plt.subplots(figsize=(10, 6))
        plot_func(ax)

        img_bytes = io.BytesIO()
        plt.tight_layout()
        fig.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
        img_bytes.seek(0)
        plt.close(fig)
        return img_bytes

    def style_combo_chart(
            ax,
            df,
            title,
            x_col,
            bar1_col,
            bar2_col,
            line_col,
            bar1_label,
            bar2_label,
            line_label,
            is_percent_line=True,
            y1_label="",
            y2_label="",
            sort_col=None,
            ascending=False
    ):
        """
        绘制类似Excel组合图：
        蓝色柱形 + 橙色柱形 + 灰色折线 + 右侧坐标轴 + 数据标签
        """

        if df is None or df.empty:
            ax.text(
                0.5, 0.5, "暂无数据",
                ha="center", va="center",
                fontsize=12, color="#666666"
            )
            ax.set_title(title, fontsize=14, color="#595959")
            ax.axis("off")
            return

        plot_df = df.copy()

        # 清洗分公司名称
        plot_df[x_col] = plot_df[x_col].astype(str).str.strip()
        plot_df = plot_df[
            ~plot_df[x_col].isin(["", "/", "nan", "None", "总计", "合计"])
        ]

        # 数值转换
        for col in [bar1_col, bar2_col, line_col]:
            plot_df[col] = pd.to_numeric(plot_df[col], errors="coerce")

        plot_df = plot_df.dropna(subset=[bar1_col, bar2_col, line_col])

        if plot_df.empty:
            ax.text(
                0.5, 0.5, "暂无有效数据",
                ha="center", va="center",
                fontsize=12, color="#666666"
            )
            ax.set_title(title, fontsize=14, color="#595959")
            ax.axis("off")
            return

            # 排序：出勤率/成绩由高到低
        if sort_col is not None and sort_col in plot_df.columns:
            plot_df = plot_df.sort_values(
                by=sort_col,
                ascending=ascending
            ).reset_index(drop=True)

        # 如果出勤率是 0~1，转换为 0~100
        if is_percent_line and plot_df[line_col].max() <= 1.5:
            plot_df[line_col] = plot_df[line_col] * 100

        x_labels = plot_df[x_col].tolist()
        x = np.arange(len(x_labels))
        width = 0.34

        # 颜色，尽量贴近你截图里的Excel风格
        blue = "#5B9BD5"
        orange = "#ED7D31"
        gray = "#A5A5A5"
        grid_gray = "#D9D9D9"
        text_gray = "#595959"

        # 柱形图
        bars1 = ax.bar(
            x - width / 2,
            plot_df[bar1_col],
            width,
            label=bar1_label,
            color=blue
        )

        bars2 = ax.bar(
            x + width / 2,
            plot_df[bar2_col],
            width,
            label=bar2_label,
            color=orange
        )

        # 左轴样式
        ax.set_title(title, fontsize=14, color=text_gray, pad=14)
        ax.set_ylabel(y1_label, fontsize=9, color=text_gray)
        ax.set_xticks(x)
        ax.set_xticklabels(
            x_labels,
            rotation=45,
            ha="right",
            fontsize=8,
            color=text_gray
        )

        ax.tick_params(axis="y", labelsize=8, colors=text_gray)
        ax.grid(axis="y", color=grid_gray, linewidth=0.8)
        ax.set_axisbelow(True)

        # 左轴范围
        left_max = max(
            plot_df[bar1_col].max(),
            plot_df[bar2_col].max()
        )
        if pd.notna(left_max) and left_max > 0:
            ax.set_ylim(0, left_max * 1.25)

        # 去掉边框，更像Excel截图
        for spine in ["top", "right", "left", "bottom"]:
            ax.spines[spine].set_visible(False)

        # 右轴折线
        ax2 = ax.twinx()

        line = ax2.plot(
            x,
            plot_df[line_col],
            label=line_label,
            color=gray,
            marker="o",
            markersize=4,
            linewidth=2.2
        )

        ax2.set_ylabel(y2_label, fontsize=9, color=text_gray)
        ax2.tick_params(axis="y", labelsize=8, colors=text_gray)

        for spine in ["top", "right", "left", "bottom"]:
            ax2.spines[spine].set_visible(False)

        # 右轴范围
        if is_percent_line:
            max_rate = plot_df[line_col].max()
            upper = max(120, int(np.ceil(max_rate / 20) * 20))
            ax2.set_ylim(0, upper)
            ticks = np.arange(0, upper + 1, 20)
            ax2.set_yticks(ticks)
            ax2.set_yticklabels([f"{int(t)}%" for t in ticks])
        else:
            max_score = plot_df[line_col].max()
            upper = max(100, max_score * 1.15)
            ax2.set_ylim(0, upper)

        # 折线数据标签
        for i, val in enumerate(plot_df[line_col]):
            if pd.notna(val):
                if is_percent_line:
                    label = f"{val:.0f}%"
                    offset = 2
                else:
                    label = f"{val:.1f}"
                    offset = 1.5

                ax2.text(
                    x[i],
                    val + offset,
                    label,
                    ha="center",
                    va="bottom",
                    fontsize=8,
                    color=text_gray
                )

        # 图例放到底部居中
        handles1, labels1 = ax.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()

        ax.legend(
            handles1 + handles2,
            labels1 + labels2,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.25),
            ncol=3,
            frameon=False,
            fontsize=8
        )

        ax.margins(x=0.01)

    def style_score_chart(
            ax,
            df,
            title,
            x_col="分公司",
            bar_col="当月实际出勤",
            line_col="平均成绩",
            bar_label="求和项：当月实际出勤",
            line_label="平均成绩",
            right_axis_max=None,
            right_tick_step=None
    ):
        """
        成绩图样式：
        - 蓝色柱形：当月实际出勤
        - 橙色折线：平均成绩
        - 左轴：平均成绩，0~120
        - 右轴：实际出勤人数
        - 按平均成绩由高到低排序
        """

        if df is None or df.empty:
            ax.text(
                0.5, 0.5, "暂无数据",
                ha="center", va="center",
                fontsize=12, color="#666666"
            )
            ax.set_title(title, fontsize=14, color="#595959")
            ax.axis("off")
            return

        plot_df = df.copy()

        # 清洗分公司名称
        plot_df[x_col] = plot_df[x_col].astype(str).str.strip()
        plot_df = plot_df[
            ~plot_df[x_col].isin(["", "/", "nan", "None", "总计", "合计"])
        ]

        # 数值转换
        plot_df[bar_col] = pd.to_numeric(plot_df[bar_col], errors="coerce")
        plot_df[line_col] = pd.to_numeric(plot_df[line_col], errors="coerce")

        plot_df = plot_df.dropna(subset=[bar_col, line_col])

        if plot_df.empty:
            ax.text(
                0.5, 0.5, "暂无有效数据",
                ha="center", va="center",
                fontsize=12, color="#666666"
            )
            ax.set_title(title, fontsize=14, color="#595959")
            ax.axis("off")
            return

        # 关键：成绩由高到低排序
        plot_df = plot_df.sort_values(
            by=line_col,
            ascending=False
        ).reset_index(drop=True)

        x_labels = plot_df[x_col].tolist()
        x = np.arange(len(x_labels))

        blue = "#5B9BD5"
        orange = "#ED7D31"
        grid_gray = "#D9D9D9"
        text_gray = "#595959"

        # 左轴：平均成绩折线
        line = ax.plot(
            x,
            plot_df[line_col],
            color=orange,
            marker="o",
            markersize=4,
            linewidth=2.4,
            label=line_label,
            zorder=3
        )

        ax.set_title(title, fontsize=14, color=text_gray, pad=14)
        ax.set_ylim(0, 120)
        ax.set_yticks(np.arange(0, 121, 20))
        ax.tick_params(axis="y", labelsize=8, colors=text_gray)
        ax.grid(axis="y", color=grid_gray, linewidth=0.8)
        ax.set_axisbelow(True)

        ax.set_xticks(x)
        ax.set_xticklabels(
            x_labels,
            rotation=45,
            ha="right",
            fontsize=8,
            color=text_gray
        )

        # 折线数据标签
        for i, val in enumerate(plot_df[line_col]):
            ax.text(
                x[i],
                val + 3,
                f"{val:.0f}",
                ha="center",
                va="bottom",
                fontsize=8,
                color=text_gray
            )

        # 右轴：实际出勤人数柱形
        ax2 = ax.twinx()

        ax.set_zorder(ax2.get_zorder() + 1)
        ax.patch.set_visible(False)

        bars = ax2.bar(
            x,
            plot_df[bar_col],
            width=0.32,
            color=blue,
            label=bar_label,
            zorder=2
        )

        # 右侧人数轴：支持渠道图按 0、20、40、60、80、100、120 这种稀疏刻度显示
        max_attend = plot_df[bar_col].max()

        if right_axis_max is not None:
            right_upper = right_axis_max
        else:
            # 自动模式：根据最大人数选择合适上限
            if max_attend <= 16:
                right_upper = 16
            elif max_attend <= 40:
                right_upper = 40
            elif max_attend <= 60:
                right_upper = 60
            elif max_attend <= 80:
                right_upper = 80
            elif max_attend <= 100:
                right_upper = 100
            else:
                right_upper = int(np.ceil(max_attend / 20) * 20)

        if right_tick_step is not None:
            tick_step = right_tick_step
        else:
            # 自动模式：人数少用 2，人数多用 10 或 20，避免刻度过密
            if right_upper <= 16:
                tick_step = 2
            elif right_upper <= 40:
                tick_step = 5
            elif right_upper <= 80:
                tick_step = 10
            else:
                tick_step = 20

        ax2.set_ylim(0, right_upper)
        ax2.set_yticks(np.arange(0, right_upper + 1, tick_step))
        ax2.tick_params(axis="y", labelsize=8, colors=text_gray)

        # 去掉边框
        for spine in ["top", "right", "left", "bottom"]:
            ax.spines[spine].set_visible(False)
            ax2.spines[spine].set_visible(False)

        # 图例：底部居中，顺序和截图一致：蓝柱在前，橙线在后
        handles1, labels1 = ax.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()

        ax.legend(
            handles2 + handles1,
            labels2 + labels1,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.23),
            ncol=2,
            frameon=False,
            fontsize=8
        )

        ax.margins(x=0.01)
    # 4. 定义4个图表的绘制函数（复用原有统计逻辑）
    # -------------------------- 内部出勤率图 --------------------------
    def plot_internal_attendance(ax):
            """内部员工出勤率图：直接使用附录Excel分公司统计sheet中的内部表。"""

            chart_df = branch_internal_summary_df.rename(columns={
                "应出勤人数": "当月应出勤",
                "实际出勤人数": "当月实际出勤"
            }).copy()

            style_combo_chart(
                ax=ax,
                df=chart_df,
                title=f"{title_prefix}IVD内部技术交流会出勤率",
                x_col="分公司",
                bar1_col="当月应出勤",
                bar2_col="当月实际出勤",
                line_col="出勤率",
                bar1_label="求和项：当月应出勤",
                bar2_label="求和项：当月实际出勤",
                line_label="出勤率",
                is_percent_line=True,
                y1_label="人数",
                y2_label="出勤率",
                sort_col="出勤率",
                ascending=False
            )



    # -------------------------- 渠道出勤率图 --------------------------
    def plot_channel_attendance(ax):
        """渠道员工出勤率图：直接使用附录Excel分公司统计sheet中的渠道表。"""

        chart_df = branch_channel_summary_df.rename(columns={
            "应出勤人数": "当月应出勤",
            "实际出勤人数": "当月实际出勤"
        }).copy()

        style_combo_chart(
            ax=ax,
            df=chart_df,
            title=f"{title_prefix}IVD渠道技术交流会出勤率",
            x_col="分公司",
            bar1_col="当月应出勤",
            bar2_col="当月实际出勤",
            line_col="出勤率",
            bar1_label="求和项：当月应出勤",
            bar2_label="求和项：当月实际出勤",
            line_label="出勤率",
            is_percent_line=True,
            y1_label="人数",
            y2_label="出勤率",
            sort_col="出勤率",
            ascending=False
        )

    # -------------------------- 内部成绩图 --------------------------
    def plot_internal_score(ax):
        """内部员工成绩图：直接使用附录Excel分公司统计sheet中的内部平均成绩。"""

        chart_df = branch_internal_summary_df.rename(columns={
            "实际出勤人数": "当月实际出勤"
        })[["分公司", "当月实际出勤", "平均成绩"]].copy()

        style_score_chart(
            ax=ax,
            df=chart_df,
            title=f"{title_prefix}IVD技术交流会-内部考核成绩",
            x_col="分公司",
            bar_col="当月实际出勤",
            line_col="平均成绩",
            bar_label="求和项：当月实际出勤",
            line_label="平均成绩"
        )

    # -------------------------- 渠道成绩图 --------------------------
    def plot_channel_score(ax):
        """渠道员工成绩图：直接使用附录Excel分公司统计sheet中的渠道平均成绩。"""

        chart_df = branch_channel_summary_df.rename(columns={
            "实际出勤人数": "当月实际出勤"
        })[["分公司", "当月实际出勤", "平均成绩"]].copy()

        style_score_chart(
            ax=ax,
            df=chart_df,
            title=f"{title_prefix}IVD技术交流会-渠道考核成绩",
            x_col="分公司",
            bar_col="当月实际出勤",
            line_col="平均成绩",
            bar_label="求和项: 当月实际出勤",
            line_label="平均成绩",
            right_axis_max=120,
            right_tick_step=20
        )

    # 5. 预生成所有图表（避免多次重复计算）
    charts = {
        "内部出勤率": generate_chart_to_bytesio(plot_internal_attendance),
        "渠道出勤率": generate_chart_to_bytesio(plot_channel_attendance),
        "内部成绩": generate_chart_to_bytesio(plot_internal_score),
        "渠道成绩": generate_chart_to_bytesio(plot_channel_score)
    }

    # 6. 精准定位并插入图表（兼容所有python-docx版本）
    insert_positions = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "具体内部出勤率如下：":
            insert_positions.append((i + 1, charts["内部出勤率"]))
        elif text == "具体渠道出勤率如下：":
            insert_positions.append((i + 1, charts["渠道出勤率"]))
        elif text == "具体内部交流会考试成绩如下：":
            insert_positions.append((i + 1, charts["内部成绩"]))
        elif text == "具体渠道交流会考试成绩如下：":
            insert_positions.append((i + 1, charts["渠道成绩"]))

    # 倒序插入（从后往前插，不影响前面的索引）
    for pos, img_bytes in reversed(insert_positions):
        # 插入空段落
        doc.paragraphs.insert(pos, "")
        para = doc.paragraphs[pos]

        # 设置段落属性
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(img_bytes, width=Inches(6))

        # 插入空行分隔
        doc.paragraphs.insert(pos + 1, "")

    # 7. 保存最终文档
    doc.save(output_path)
    print(f"总结文档已生成并保存至: {output_path}")


def generate_dept_avg_chart(
    merged_data: pd.DataFrame,
    staff_type: str,
    save_path: str,
    title: str = "分公司平均成绩"
):
    """生成类似截图风格的分公司平均成绩组合图"""

    exam_types = ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]

    staff_data = merged_data[merged_data["人员类型"] == staff_type].copy()

    dept_avg = []
    for dept, group in staff_data.groupby("分公司"):
        dept_scores = []

        for _, row in group.iterrows():
            present_subjects = _get_row_present_subjects(row)

            if not present_subjects:
                continue

            dept_scores.append(_calculate_average_from_present_subjects(row, exam_types))

        if dept_scores:
            dept_avg.append({
                "分公司": dept,
                "满分": 100,
                "平均成绩": round(sum(dept_scores) / len(dept_scores), 1),
                "平均成绩折线": round(sum(dept_scores) / len(dept_scores), 1)
            })

    dept_avg_df = pd.DataFrame(dept_avg)

    if dept_avg_df.empty:
        print(f"⚠️ {staff_type}平均成绩数据为空，跳过绘图")
        return

    _plot_combo_bar_line(
        df=dept_avg_df,
        branch_col="分公司",
        bar1_col="满分",
        bar2_col="平均成绩",
        line_col="平均成绩折线",
        save_path=save_path,
        title=title,
        bar1_label="满分",
        bar2_label="平均成绩",
        line_label="平均成绩",
        y1_label="成绩",
        y2_label="成绩",
        line_is_percent=False
    )


def generate_attendance_chart(
    merged_data: pd.DataFrame,
    staff_type: str,
    save_path: str,
    title: str = "分公司出勤率"
):
    """生成类似截图风格的分公司出勤率组合图"""

    staff_data = merged_data[merged_data["人员类型"] == staff_type].copy()

    attendance_data = []
    for dept, group in staff_data.groupby("分公司"):
        should_attend = len(group)
        actual_attend = int(group["attended"].sum())
        attendance_rate = actual_attend / should_attend * 100 if should_attend > 0 else 0

        attendance_data.append({
            "分公司": dept,
            "当月应出勤": should_attend,
            "当月实际出勤": actual_attend,
            "出勤率": attendance_rate
        })

    attendance_df = pd.DataFrame(attendance_data)

    if attendance_df.empty:
        print(f"⚠️ {staff_type}出勤率数据为空，跳过绘图")
        return

    _plot_combo_bar_line(
        df=attendance_df,
        branch_col="分公司",
        bar1_col="当月应出勤",
        bar2_col="当月实际出勤",
        line_col="出勤率",
        save_path=save_path,
        title=title,
        bar1_label="求和项：当月应出勤",
        bar2_label="求和项：当月实际出勤",
        line_label="出勤率",
        y1_label="人数",
        y2_label="出勤率",
        line_is_percent=True
    )


# -------------------------- 主执行函数（整合所有逻辑） --------------------------
def main(
        roster_file: str,  # 应出勤名单文件路径
        凝血_file: str,  # 凝血成绩文件
        化免_file: str,  # 化免成绩文件
        临检_file: str,  # 临检成绩文件
        TLA_file: str,  # TLA成绩文件
        微生物_file: str,  # 微生物成绩文件
        urine_file: str,  # 尿液成绩文件
        output_path: str = "交流会统计结果.xlsx"  # 输出文件路径（默认值）
):
    # 1. 加载应出勤名单
    roster_df = load_attendance_roster(roster_file)

    # 2. 构造各科成绩文件字典
    exam_files = {
        '凝血': 凝血_file,
        '化免': 化免_file,
        '临检': 临检_file,
        'TLA': TLA_file,
        '微生物': 微生物_file,
        '尿液': urine_file
    }

    # 3. 加载所有科目成绩
    all_exam_dfs = []
    for exam_type, file_path in exam_files.items():
        try:
            exam_df = load_exam_scores(exam_type, file_path)
            all_exam_dfs.append(exam_df)
        except Exception as e:
            print(f"加载{exam_type}成绩失败: {str(e)}")

    # 4. 关联名单和成绩数据
    merged_data = merge_attendance_with_scores(roster_df, all_exam_dfs)

    # 5. 生成主统计文件
    create_main_statistics_file(roster_file, exam_files, output_path, merged_data)

    # （可选）生成附录文件
    absentee_lists = generate_absentee_lists(merged_data)
    
    create_appendix_workbook(merged_data, absentee_lists, roster_df, "中国区技术交流会出勤和考核情况-IVD线.xlsx")

    overall_info=calculate_overall_summary(merged_data)
    meeting_month_label = _get_meeting_month_label_from_exam_files(exam_files)

    generate_summary_docx(
        merged_data,
        absentee_lists,
        overall_info,
        roster_df,
        meeting_month_label=meeting_month_label,
        branch_summary_excel_path="中国区技术交流会出勤和考核情况-IVD线.xlsx"
    )

    print("统计文件生成完成！")


# 主程序调用示例（放在代码末尾）

def main(
    凝血_file: str,
    化免_file: str,
    临检_file: str,
    TLA_file: str,
    微生物_file: str,
    urine_file: str,
    output_path: str = "交流会统计结果.xlsx",
    template_path: str = "中国区技术交流会总结邮件内容模版.docx",
    word_output_path: str = "中国区技术交流会总结邮件.docx"
):
    """
    技术交流会统计主函数。

    新规则：
    1. 不再输入名单.xlsx；
    2. 名单由六科成绩中所有账号去重汇总；
    3. 渠道人员解析渠道商；
    4. 输出：
       交流会统计结果.xlsx
       中国区技术交流会出勤和考核情况-IVD线.xlsx
       Word 总结
    """

    print("\n" + "=" * 80)
    print("开始执行中国区技术交流会统计...")
    print(f"凝血成绩：{凝血_file}")
    print(f"化免成绩：{化免_file}")
    print(f"临检成绩：{临检_file}")
    print(f"TLA成绩：{TLA_file}")
    print(f"微生物成绩：{微生物_file}")
    print(f"尿液成绩：{urine_file}")
    print(f"Excel输出：{output_path}")
    print(f"Word模板：{template_path}")
    print(f"Word输出：{word_output_path}")

    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    word_output_dir = os.path.dirname(word_output_path)
    if word_output_dir:
        os.makedirs(word_output_dir, exist_ok=True)

    exam_files = {
        "凝血": 凝血_file,
        "化免": 化免_file,
        "临检": 临检_file,
        "TLA": TLA_file,
        "微生物": 微生物_file,
        "尿液": urine_file
    }

    all_exam_dfs = []

    for exam_type, file_path in exam_files.items():
        print(f"正在读取 {exam_type} 成绩文件：{file_path}")
        exam_df = load_exam_scores(exam_type, file_path)
        all_exam_dfs.append(exam_df)

    # 新规则：从六科成绩中汇总名单
    roster_df = build_roster_from_exam_dfs(all_exam_dfs)

    print(f"名单汇总完成：共 {len(roster_df)} 人")

    merged_data = merge_attendance_with_scores(
        roster_df,
        all_exam_dfs
    )

    # 关键补丁：从六科成绩表回填分公司、渠道商等人员信息
    merged_data = enrich_person_info_from_exam_dfs(
        merged_data,
        all_exam_dfs
    )

    # 新出勤口径：所有成绩均为 0 或没有成绩的人算不出勤；否则算实际出勤。
    merged_data = _apply_attendance_by_score_logic(
        merged_data,
        ["凝血", "化免", "临检", "TLA", "微生物", "尿液"]
    )

    overall_summary = calculate_overall_summary(merged_data)

    absentee_lists = generate_absentee_lists(merged_data)

    print(f"正在生成主统计 Excel：{output_path}")

    create_main_statistics_file(
        roster_file="",
        exam_files=exam_files,
        output_path=output_path,
        merged_data=merged_data
    )

    output_dir = os.path.dirname(output_path)
    output_filename = os.path.basename(output_path)

    if not output_dir:
        output_dir = os.getcwd()

    appendix_output_path = os.path.join(
        output_dir,
        "中国区技术交流会出勤和考核情况-IVD线.xlsx"
    )

    print(f"正在生成附录 Excel：{appendix_output_path}")

    create_appendix_workbook(
        merged_data,
        absentee_lists,
        roster_df,
        appendix_output_path
    )

    print(f"正在生成 Word 总结：{word_output_path}")

    meeting_month_label = _get_meeting_month_label_from_exam_files(exam_files)

    generate_summary_docx(
        merged_data=merged_data,
        absentee_lists=absentee_lists,
        overall_summary=overall_summary,
        roster_df=roster_df,
        template_path=template_path,
        output_path=word_output_path,
        meeting_month_label=meeting_month_label,
        branch_summary_excel_path=appendix_output_path
    )

    print("交流会统计处理完成！")
    print("=" * 80)

    return {
        "merged_data": merged_data,
        "overall_summary": overall_summary,
        "absentee_lists": absentee_lists,
        "main_excel": output_path,
        "appendix_excel": appendix_output_path,
        "word_report": word_output_path
    }

