# -*- coding: utf-8 -*-
"""
big_teach.py
大练兵成绩统计底层逻辑（可选补考、动态科目、可选Word版）

本版本兼容三种输入场景：
1. 只上传第一次考试成绩；
2. 上传第一次考试成绩 + 第一次补考成绩；
3. 上传第一次考试成绩 + 第一次补考成绩 + 第二次补考/重修成绩；
4. Word模板可不上传，只生成Excel。

对外函数保持不变：
- merge_all_exam_data(first_folder, remake_folder='', resit_folder='')
- generate_summary_excel(merged_data, output_path)
- generate_analysis_report(excel_path, template_path, output_path)
"""

import os
import re
import warnings
from typing import Dict, List, Tuple

import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openpyxl.styles import Font, Alignment

warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')
warnings.filterwarnings('ignore', category=FutureWarning)


SUBJECT_LIST = ['流式', '临检基础医疗', '临检产线', '微生物', '尿液', '化免产线', '化免基础医疗', '凝血', '免疫', '生化', 'TLA', '临检JC&社会办医']
SUBJECT_LIST_SORTED = sorted(SUBJECT_LIST, key=len, reverse=True)

COLUMN_KEYWORD_MAP = {
    '账号': ['账号', '工号', '编号', '用户账号', '员工编号', '考生号', '学号', '员工工号'],
    '姓名': ['姓名', '名字', '员工姓名', '考生姓名', '用户名', '真实姓名'],
    '部门': ['部门', '所属部门', '组织架构', '单位', '所属单位', '组织机构'],
    '岗位': ['岗位', '职位', '职务', '角色', '职称'],
    '成绩': ['成绩', '得分', '分数', '考试得分', '最终得分', '卷面分', '考试分数']
}

BRANCH_TO_AREA = {
    '北京分公司': '华北区', '天津分公司': '华北区', '石家庄分公司': '华北区',
    '青岛分公司': '华北区', '济南分公司': '华北区', '北京战略客户服务部': '华北区',
    '上海分公司': '华东区', '南京分公司': '华东区', '杭州分公司': '华东区',
    '武汉分公司': '华东区', '苏州分公司': '华东区', '合肥分公司': '华东区',
    '广州分公司': '华南区', '深圳分公司': '华南区', '长沙分公司': '华南区',
    '南昌分公司': '华南区', '海南分公司': '华南区', '福州分公司': '华南区',
    '深圳办事处': '华南区',
    '成都分公司': '西南区', '重庆分公司': '西南区', '南宁分公司': '西南区',
    '贵阳分公司': '西南区', '昆明分公司': '西南区',
    '西安分公司': '西北区', '兰州分公司': '西北区', '郑州分公司': '西北区',
    '新疆分公司': '西北区', '太原分公司': '西北区',
    '沈阳分公司': '东北区', '内蒙古分公司': '东北区', '黑吉分公司': '东北区'
}

# 分公司统计合并口径：这些部门/办事处在“分公司分析”和Word分公司表中并入对应分公司。
BRANCH_ALIAS_MAP = {
    '深圳办事处': '深圳分公司',
    '北京战略客户服务部': '北京分公司',
    '北京战略服务部': '北京分公司',
}

AREA_ORDER = ['华东区', '华南区', '西南区', '西北区', '东北区', '华北区']

DRILL_SEASONS = ['春季', '夏季', '秋季', '冬季']
DEFAULT_DRILL_SEASON = '春季'


def _extract_drill_season_from_text(text) -> str:
    """从文件名/文本中识别春季、夏季、秋季、冬季大练兵；识别不到返回空字符串。"""
    text = '' if text is None else str(text)
    for season in DRILL_SEASONS:
        if season in text and '大练兵' in text:
            return season
    for season in DRILL_SEASONS:
        if f'{season}大练兵' in text or f'IVD{season}' in text:
            return season
    return ''


def _normalize_drill_season(season: str) -> str:
    season = str(season or '').strip()
    return season if season in DRILL_SEASONS else DEFAULT_DRILL_SEASON


def infer_drill_season_from_folders(*folder_paths) -> str:
    """从已选择的成绩文件夹文件名中识别大练兵季节，供GUI提前生成输出文件名。"""
    seasons = []
    for folder_path in folder_paths:
        if _is_optional_folder_empty(folder_path):
            continue
        folder_path = str(folder_path).strip()
        if not os.path.isdir(folder_path):
            continue
        for file_path in _iter_excel_files(folder_path):
            season = _extract_drill_season_from_text(os.path.basename(file_path))
            if season:
                seasons.append(season)

    for season in DRILL_SEASONS:
        if season in seasons:
            return season
    return DEFAULT_DRILL_SEASON


def get_drill_season_from_merged_data(merged_data) -> str:
    """从 merge_all_exam_data 返回的数据中读取本次大练兵季节。"""
    if isinstance(merged_data, dict):
        meta = merged_data.get('__meta__')
        if isinstance(meta, dict):
            season = meta.get('大练兵季节', '')
            if season in DRILL_SEASONS:
                return season
        for value in merged_data.values():
            if isinstance(value, pd.DataFrame) and '大练兵季节' in value.columns and not value.empty:
                values = value['大练兵季节'].dropna().astype(str).tolist()
                for season in DRILL_SEASONS:
                    if season in values:
                        return season
    return DEFAULT_DRILL_SEASON


def get_drill_season_from_path(*paths) -> str:
    text = ' '.join(str(p or '') for p in paths)
    season = _extract_drill_season_from_text(text)
    return _normalize_drill_season(season)


def get_drill_output_filenames(season: str) -> Tuple[str, str]:
    season = _normalize_drill_season(season)
    return (
        f'IVD{season}大练兵成绩汇总.xlsx',
        f'IVD用户服务工程师{season}大练兵分析报告.docx'
    )


def _replace_drill_season_in_text(text: str, season: str) -> str:
    """将文本中的春/夏/秋/冬季大练兵统一替换为本次识别到的季节大练兵。"""
    season = _normalize_drill_season(season)
    text = '' if text is None else str(text)
    text = re.sub(r'(春季|夏季|秋季|冬季)大练兵', f'{season}大练兵', text)
    return text


def _replace_drill_season_in_doc(doc, season: str):
    season = _normalize_drill_season(season)

    def replace_runs(paragraph):
        for run in paragraph.runs:
            if run.text:
                run.text = _replace_drill_season_in_text(run.text, season)

    for paragraph in doc.paragraphs:
        replace_runs(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_runs(paragraph)


def _clean_text(value) -> str:
    if pd.isna(value):
        return ''
    return str(value).strip()


def _normalize_account(value) -> str:
    text = _clean_text(value)
    if text.endswith('.0'):
        text = text[:-2]
    return text.strip()


def _normalize_branch_name(branch) -> str:
    """统一分公司名称，用于分公司分析和Word分公司表统计。"""
    branch_text = _clean_text(branch)
    return BRANCH_ALIAS_MAP.get(branch_text, branch_text)


def _is_blank_for_anomaly(value) -> bool:
    text = _clean_text(value)
    return text in ['', '/', 'nan', 'None', '<NA>', '未知分公司', '未知区域']


def _build_post_remark(post) -> str:
    """备注列口径：除“工程师”“用户服务工程师”外，其余岗位均显示岗位名称。"""
    post_text = _clean_text(post)
    if post_text in ['', 'nan', 'None', '<NA>']:
        return ''
    if post_text in ['工程师', '用户服务工程师']:
        return ''
    return post_text


def _to_numeric_score(value) -> float:
    if pd.isna(value):
        return 0.0
    text = str(value).strip()
    if text in ['', '--', '-', '/', '未参考', '未提交', 'nan', 'None']:
        return 0.0
    try:
        return float(text)
    except Exception:
        nums = re.findall(r'\d+(?:\.\d+)?', text)
        return float(nums[-1]) if nums else 0.0


def _rename_columns_fuzzy(df: pd.DataFrame) -> pd.DataFrame:
    column_mapping = {}
    for standard_col, keywords in COLUMN_KEYWORD_MAP.items():
        for col in df.columns:
            col_text = str(col).strip()
            if any(keyword in col_text for keyword in keywords):
                column_mapping[col] = standard_col
                break
    return df.rename(columns=column_mapping)


def _is_optional_folder_empty(folder_path) -> bool:
    return folder_path is None or str(folder_path).strip() == ''


def _iter_excel_files(folder_path: str) -> List[str]:
    excel_files = []
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            lower = file_name.lower()
            if lower.startswith('~$'):
                continue
            if lower.endswith('.xlsx'):
                excel_files.append(os.path.join(root, file_name))
    return excel_files


def _is_resit_file_name(file_name: str) -> bool:
    """识别第一次补考文件名，兼容【补考】、（补考）、(补考)等写法。"""
    name = str(file_name)
    # 第二次补考/重修优先归入 remake，避免被普通“补考”误判为第一次补考。
    if _is_remake_file_name(name):
        return False
    return any(keyword in name for keyword in [
        '【补考】', '（补考）', '(补考)', '补考'
    ])


def _is_remake_file_name(file_name: str) -> bool:
    """识别第二次补考/重修文件名。"""
    name = str(file_name)
    return any(keyword in name for keyword in [
        '【重修】', '（重修）', '(重修)', '重修',
        '第二次补考', '二次补考', '第2次补考', '2次补考'
    ])


def _is_makeup_exam_file_name(file_name: str) -> bool:
    """识别任意补考/重修文件，用于第一次考试文件夹排除补考文件。"""
    return _is_resit_file_name(file_name) or _is_remake_file_name(file_name)


def read_exam_data(folder_path, exam_type: str) -> Dict[str, pd.DataFrame]:
    """
    读取考试数据。

    exam_type:
    - first  第一次考试
    - resit  第一次补考
    - remake 第二次补考/重修

    本函数已支持空文件夹路径：
    如果 resit/remake 未选择，直接返回空字典，不影响后续统计。
    """
    exam_data = {}

    if _is_optional_folder_empty(folder_path):
        print(f"ℹ️ {exam_type} 文件夹未选择，跳过读取。")
        return exam_data

    folder_path = str(folder_path).strip()

    if not os.path.exists(folder_path):
        print(f"⚠️ {exam_type} 文件夹不存在，已跳过：{folder_path}")
        return exam_data

    # 补考/重修文件名兼容多种括号写法：
    # 第一次补考：支持 “【补考】”、“（补考）”、“(补考)” 和普通“补考”。
    # 第二次补考/重修：支持 “【重修】”、“（重修）”、“(重修)”、“重修”、“第二次补考”等。
    # 第一次考试读取时会自动排除这些补考/重修文件。

    header_keywords = ['账号', '工号', '姓名', '部门', '成绩', '得分', '岗位']
    excel_files = _iter_excel_files(folder_path)

    print(f"\n📂 正在扫描目录：{folder_path}")
    print(f"✅ 共发现 Excel 文件：{len(excel_files)} 个")

    for file_path in excel_files:
        file_name = os.path.basename(file_path)

        if exam_type == 'first':
            if _is_makeup_exam_file_name(file_name):
                continue
        elif exam_type == 'resit':
            if not _is_resit_file_name(file_name):
                continue
        elif exam_type == 'remake':
            if not _is_remake_file_name(file_name):
                continue

        matched_subject = None
        for subject in SUBJECT_LIST_SORTED:
            if subject in file_name:
                matched_subject = subject
                break

        if not matched_subject:
            print(f"⚠️ 未识别科目，跳过文件：{file_name}")
            continue

        try:
            print(f"\n{'=' * 60}")
            print(f"正在读取文件：{file_name}")
            print(f"识别科目：{matched_subject} | 考试类型：{exam_type}")

            xl = pd.ExcelFile(file_path)
            sheet_name = xl.sheet_names[0]
            df_raw = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

            header_row_index = -1
            max_keyword_count = 0
            for i in range(min(15, len(df_raw))):
                row_content = str(df_raw.iloc[i].tolist()).lower()
                keyword_count = sum(1 for kw in header_keywords if kw.lower() in row_content)
                if keyword_count > max_keyword_count:
                    max_keyword_count = keyword_count
                    header_row_index = i

            if header_row_index == -1 or max_keyword_count < 3:
                print(f"❌ 未找到有效表头，跳过：{file_name}")
                continue

            df = df_raw.iloc[header_row_index + 1:].copy()
            df.columns = df_raw.iloc[header_row_index].tolist()
            df.columns = (
                df.columns.astype(str)
                .str.strip()
                .str.replace('\n', '', regex=False)
                .str.replace('\r', '', regex=False)
            )
            df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')
            df = _rename_columns_fuzzy(df)

            if '账号' in df.columns:
                df['工号'] = df['账号'].apply(_normalize_account)
            elif '工号' in df.columns:
                df['工号'] = df['工号'].apply(_normalize_account)
            else:
                print(f"❌ 缺少账号/工号列，跳过：{file_name}")
                continue

            if '成绩' not in df.columns:
                print(f"❌ 缺少成绩列，跳过：{file_name}")
                continue

            for col in ['姓名', '部门', '岗位']:
                if col not in df.columns:
                    df[col] = ''

            df['考试类型'] = exam_type
            df['考试科目'] = matched_subject
            df['大练兵季节'] = _extract_drill_season_from_text(file_name)

            if matched_subject in exam_data:
                exam_data[matched_subject] = pd.concat([exam_data[matched_subject], df], ignore_index=True)
            else:
                exam_data[matched_subject] = df

            print(f"✅ 文件读取成功：{matched_subject} | {len(df)} 行")

        except Exception as e:
            print(f"❌ 读取失败：{file_name} | {e}")
            import traceback
            traceback.print_exc()

    if not exam_data:
        print(f"ℹ️ {exam_type} 未读取到有效考试文件。")

    return exam_data


def _parse_department(dept, emp_type: str) -> Tuple[str, str, str]:
    dept = _clean_text(dept)
    if not dept:
        return '未知区域', '未知分公司', '/' if emp_type == '内部' else ''

    text = (dept.replace('&gt;', '>')
                .replace('－>', '->')
                .replace('—>', '->')
                .replace('→', '->')
                .replace('＞', '>'))
    parts = re.split(r'\s*[-－—=]*>\s*', text)
    parts = [p.strip() for p in parts if p and p.strip()]

    area = '未知区域'
    branch = '未知分公司'
    channel_name = '/' if emp_type == '内部' else ''

    if emp_type == '渠道':
        for part in parts:
            if '分公司' in part or '办事处' in part or '战略客户服务部' in part or '战略服务部' in part:
                branch = part
                break
        if branch in parts:
            idx = parts.index(branch)
            if idx + 1 < len(parts):
                channel_name = parts[idx + 1]
        elif len(parts) >= 3:
            channel_name = parts[-1]
    else:
        for part in reversed(parts):
            if '分公司' in part or '办事处' in part or '战略客户服务部' in part or '战略服务部' in part or '服务部' in part:
                branch = part
                break
        if branch == '未知分公司' and parts:
            branch = parts[-1]

    branch = _normalize_branch_name(branch)
    area = BRANCH_TO_AREA.get(branch, '未知区域')
    return area, branch, channel_name


def merge_all_exam_data(first_folder, remake_folder='', resit_folder=''):
    """
    合并大练兵三类考试数据。

    兼容三种输入：
    - 只传 first_folder；
    - 传 first_folder + resit_folder；
    - 传 first_folder + resit_folder + remake_folder。

    参数顺序保持旧版不变：merge_all_exam_data(first_folder, remake_folder, resit_folder)
    """
    unknown_area_records = []

    def read_and_process_all_exams():
        all_exam_data = []
        exam_types = [
            (first_folder, 'first'),
            (remake_folder, 'remake'),
            (resit_folder, 'resit')
        ]

        for folder, exam_type in exam_types:
            exam_data = read_exam_data(folder, exam_type)
            for subject, df in exam_data.items():
                df = _rename_columns_fuzzy(df.copy())

                if '工号' not in df.columns:
                    if '账号' in df.columns:
                        df['工号'] = df['账号'].apply(_normalize_account)
                    else:
                        continue

                for col in ['姓名', '部门', '岗位', '成绩']:
                    if col not in df.columns:
                        df[col] = '' if col != '成绩' else 0

                df['工号'] = df['工号'].apply(_normalize_account)
                df = df[df['工号'] != ''].copy()
                df['考试类型'] = exam_type
                df['考试科目'] = subject
                all_exam_data.append(df)

        return pd.concat(all_exam_data, ignore_index=True) if all_exam_data else pd.DataFrame()

    print("\n" + "=" * 80)
    print("正在读取并统一所有考试数据...")
    all_raw_data = read_and_process_all_exams()

    if all_raw_data.empty:
        print("❌ 未读取到任何考试数据！")
        return {}

    season_values = []
    if '大练兵季节' in all_raw_data.columns:
        season_values = [x for x in all_raw_data['大练兵季节'].dropna().astype(str).tolist() if x in DRILL_SEASONS]
    drill_season = DEFAULT_DRILL_SEASON
    for season in DRILL_SEASONS:
        if season in season_values:
            drill_season = season
            break

    print(f"✅ 共读取到 {len(all_raw_data)} 条考试记录")
    print(f"✅ 已识别本次为：{drill_season}大练兵")

    # 过滤禁用人员
    if '姓名' in all_raw_data.columns:
        all_raw_data = all_raw_data[~all_raw_data['姓名'].astype(str).str.contains('已禁用', na=False)].copy()

    global_base_info = {}
    for _, row in all_raw_data.iterrows():
        emp_id = _normalize_account(row.get('工号', ''))
        if not emp_id:
            continue

        current_info = {
            '姓名': _clean_text(row.get('姓名', '')),
            '部门': _clean_text(row.get('部门', '')),
            '岗位': _clean_text(row.get('岗位', ''))
        }

        if emp_id not in global_base_info:
            global_base_info[emp_id] = current_info
        else:
            if current_info['部门'] and not global_base_info[emp_id].get('部门'):
                global_base_info[emp_id]['部门'] = current_info['部门']
            if len(current_info['姓名']) > len(global_base_info[emp_id].get('姓名', '')):
                global_base_info[emp_id]['姓名'] = current_info['姓名']
            if current_info['岗位'] and not global_base_info[emp_id].get('岗位'):
                global_base_info[emp_id]['岗位'] = current_info['岗位']

    merged_data = {}

    for subject in SUBJECT_LIST:
        print("\n" + "=" * 80)
        print(f"正在合并 {subject} 科目数据")

        subject_data = all_raw_data[all_raw_data['考试科目'] == subject].copy()
        if subject_data.empty:
            print(f"⚠️ {subject} 科目无数据，跳过")
            continue

        all_emp_ids = subject_data['工号'].dropna().astype(str).str.strip().unique().tolist()
        subject_base_info = []
        for emp_id in all_emp_ids:
            info = global_base_info.get(emp_id, {'姓名': '', '部门': '', '岗位': ''})
            subject_base_info.append({
                '工号': emp_id,
                '姓名': info.get('姓名', ''),
                '部门': info.get('部门', ''),
                '岗位': info.get('岗位', '')
            })

        df_base = pd.DataFrame(subject_base_info)

        def score_frame(exam_type: str, col_name: str) -> pd.DataFrame:
            temp = subject_data[subject_data['考试类型'] == exam_type][['工号', '成绩']].copy()
            if temp.empty:
                return pd.DataFrame(columns=['工号', col_name])
            temp[col_name] = temp['成绩'].apply(_to_numeric_score)
            temp = temp[['工号', col_name]]
            temp = temp.sort_values(col_name, ascending=False).drop_duplicates(subset='工号', keep='first')
            return temp

        first_scores = score_frame('first', '第一次考试成绩')
        resit_scores = score_frame('resit', '第一次补考成绩')
        remake_scores = score_frame('remake', '第二次补考成绩')

        df_merge = df_base.merge(first_scores, on='工号', how='left')
        df_merge = df_merge.merge(resit_scores, on='工号', how='left')
        df_merge = df_merge.merge(remake_scores, on='工号', how='left')

        imported_exam_types = set(subject_data['考试类型'].dropna().astype(str).tolist())
        score_col_exam_type_map = {
            '第一次考试成绩': 'first',
            '第一次补考成绩': 'resit',
            '第二次补考成绩': 'remake',
        }
        for col, exam_type_for_col in score_col_exam_type_map.items():
            if col not in df_merge.columns:
                df_merge[col] = pd.NA

            # 如果该科目本次没有导入对应考试类型，则展示为“/”；
            # 如果该科目导入了对应考试类型，但某个人没有成绩记录，则按0处理。
            if exam_type_for_col in imported_exam_types:
                df_merge[col] = pd.to_numeric(df_merge[col], errors='coerce').fillna(0.0)
            else:
                df_merge[col] = pd.NA

        df_merge['员工类型'] = df_merge['工号'].apply(lambda x: '渠道' if str(x).startswith('S') else '内部')

        parsed = df_merge.apply(lambda row: pd.Series(_parse_department(row.get('部门', ''), row.get('员工类型', '内部'))), axis=1)
        parsed.columns = ['区域', '分公司', '渠道商名称']
        df_merge[['区域', '分公司', '渠道商名称']] = parsed
        df_merge.loc[df_merge['员工类型'] == '内部', '渠道商名称'] = '/'

        # 不再剔除“用户服务主管”等岗位人员：
        # 只要原始成绩表中存在该人员，科目成绩 sheet、成绩汇总和后续分析均保留。

        # 备注列：除“工程师”“用户服务工程师”外，其余岗位均显示岗位名称。
        df_merge['备注'] = ''
        if '岗位' in df_merge.columns:
            df_merge['备注'] = df_merge['岗位'].apply(_build_post_remark)

        df_merge['考试科目'] = subject
        df_merge['大练兵季节'] = drill_season

        # 异常人员：找不到具体分公司，或渠道人员找不到渠道商时，输出到“异常人员”sheet。
        for _, row in df_merge.iterrows():
            reasons = []
            if _is_blank_for_anomaly(row.get('分公司', '')):
                reasons.append('找不到具体分公司')
            if _is_blank_for_anomaly(row.get('区域', '')):
                reasons.append('找不到区域')
            if row.get('员工类型', '') == '渠道' and _is_blank_for_anomaly(row.get('渠道商名称', '')):
                reasons.append('找不到渠道商')

            if reasons:
                unknown_area_records.append({
                    '考试科目': subject,
                    '工号': row.get('工号', ''),
                    '姓名': row.get('姓名', ''),
                    '员工类型': row.get('员工类型', ''),
                    '岗位': row.get('岗位', ''),
                    '原始部门字符串': row.get('部门', ''),
                    '解析出的分公司': row.get('分公司', ''),
                    '解析出的渠道商': row.get('渠道商名称', ''),
                    '异常原因': '；'.join(reasons)
                })

        merged_data[subject] = df_merge.reset_index(drop=True)
        print(f"✅ {subject} 科目合并完成！总人数：{len(df_merge)}")

    if unknown_area_records:
        try:
            pd.DataFrame(unknown_area_records).to_excel('未知区域人员排查清单.xlsx', index=False)
            print(f"⚠️ 已导出未知区域人员排查清单.xlsx，共 {len(unknown_area_records)} 条")
        except Exception as e:
            print(f"⚠️ 未知区域人员清单导出失败：{e}")

    merged_data['__meta__'] = {
        '大练兵季节': drill_season,
        '异常人员': unknown_area_records
    }
    return merged_data


def calculate_final_score(row) -> float:
    first_score = _to_numeric_score(row.get('第一次考试成绩', 0))
    remake_score = _to_numeric_score(row.get('第二次补考成绩', 0))
    resit_score = _to_numeric_score(row.get('第一次补考成绩', 0))

    has_retake = (remake_score > 0) or (resit_score > 0)
    max_score = max(first_score, remake_score, resit_score)

    if has_retake:
        return min(max_score, 80)
    return max_score


def calculate_personal_composite_score(personal_records: pd.DataFrame) -> float:
    """
    个人成绩汇总口径：取所有有效科目最终成绩中的最高分。

    说明：
    1. 不再按尿液/微生物等科目设置权重；
    2. 不再计算算术平均；
    3. 最终用于“成绩汇总”sheet 的“平均成绩”列，以及后续基于该列生成的分析 sheet。
    """
    if personal_records is None or personal_records.empty:
        return 0.0

    if '最终成绩' not in personal_records.columns:
        return 0.0

    scores = []
    for value in personal_records['最终成绩'].tolist():
        try:
            score = _to_numeric_score(value)
        except Exception:
            score = 0.0
        scores.append(score)

    return max(scores) if scores else 0.0


def _format_percent(value: float) -> float:
    try:
        return round(float(value), 2)
    except Exception:
        return 0.0



def _is_zero_or_empty_score(value) -> bool:
    """
    缺考判断用：0、空值、/ 都视为未取得有效成绩；只要有一个科目成绩 > 0，就不算缺考。
    """
    if pd.isna(value):
        return True
    text = str(value).strip()
    if text in ['', '/', 'nan', 'None', '<NA>', '--', '-']:
        return True
    try:
        return float(text) == 0
    except Exception:
        return True


def _has_positive_score(value) -> bool:
    """是否存在有效正分成绩。"""
    return not _is_zero_or_empty_score(value)


def _calc_employee_absent_from_group(group: pd.DataFrame) -> bool:
    """
    大练兵缺勤/缺考口径：该员工所有已导入科目的最终成绩都为0或空。
    未导入的科目视为空，不影响“是否有正分成绩”的判断。
    """
    if group is None or group.empty or '最终成绩' not in group.columns:
        return True
    return not any(_has_positive_score(v) for v in group['最终成绩'].tolist())


def _build_employee_summary(all_data: pd.DataFrame) -> pd.DataFrame:
    """生成员工级汇总，区域/分公司分析使用该口径。"""
    personal_summary = []
    if all_data is None or all_data.empty:
        return pd.DataFrame()

    for emp_id, group in all_data.groupby('工号'):
        base_info = group.iloc[0]
        is_absent = _calc_employee_absent_from_group(group)

        if is_absent:
            personal_score = 0.0
        else:
            valid_group = group.copy()
            valid_group['最终成绩'] = valid_group['最终成绩'].apply(_to_numeric_score)
            personal_score = calculate_personal_composite_score(valid_group)

        personal_summary.append({
            '工号': emp_id,
            '区域': base_info.get('区域', '未知区域'),
            '分公司': base_info.get('分公司', '未知分公司'),
            '综合成绩': personal_score,
            '员工类型': '渠道' if str(emp_id).startswith('S') else '内部',
            '是否缺勤': is_absent
        })

    return pd.DataFrame(personal_summary)


def _build_subject_dimension_analysis(all_data: pd.DataFrame) -> pd.DataFrame:
    """
    产品线/考试科目维度单独计算，保证导入了哪些科目就展示哪些科目。
    产品线只反映该科目的参考、成绩和通过情况，不展示缺勤人数/缺勤率；
    因为某产品线缺考或未导入，不等同于员工整体缺勤。
    """
    if all_data is None or all_data.empty or '考试科目' not in all_data.columns:
        return pd.DataFrame()

    emp_summary = _build_employee_summary(all_data)
    absent_map = dict(zip(emp_summary.get('工号', []), emp_summary.get('是否缺勤', []))) if not emp_summary.empty else {}
    result_rows = []

    for subject in [s for s in SUBJECT_LIST if s in set(all_data['考试科目'].dropna().astype(str))]:
        group = all_data[all_data['考试科目'] == subject].copy()
        if group.empty:
            continue

        group['员工类型'] = group['工号'].apply(lambda x: '渠道' if str(x).startswith('S') else '内部')
        group['是否缺勤'] = group['工号'].map(absent_map).fillna(False).astype(bool)
        group['最终成绩_num'] = group['最终成绩'].apply(_to_numeric_score)

        total_people = int(group['工号'].nunique())

        attended_group = group[~group['是否缺勤']].copy()
        avg_score = attended_group['最终成绩_num'].mean() if not attended_group.empty else 0.0
        pass_people = int(attended_group.loc[attended_group['最终成绩_num'] >= 80, '工号'].nunique()) if not attended_group.empty else 0
        pass_rate = pass_people / total_people * 100 if total_people else 0

        internal_group = group[group['员工类型'] == '内部'].copy()
        internal_total = int(internal_group['工号'].nunique())
        internal_pass = int(internal_group.loc[(~internal_group['是否缺勤']) & (internal_group['最终成绩_num'] >= 80), '工号'].nunique()) if not internal_group.empty else 0
        internal_pass_rate = internal_pass / internal_total * 100 if internal_total else 0

        channel_group = group[group['员工类型'] == '渠道'].copy()
        channel_total = int(channel_group['工号'].nunique())
        channel_pass = int(channel_group.loc[(~channel_group['是否缺勤']) & (channel_group['最终成绩_num'] >= 80), '工号'].nunique()) if not channel_group.empty else 0
        channel_pass_rate = channel_pass / channel_total * 100 if channel_total else 0

        result_rows.append({
            '考试科目': subject,
            '总人数': total_people,
            '平均分': round(float(avg_score), 2) if pd.notna(avg_score) else 0.0,
            '通过人数': pass_people,
            '通过率(%)': _format_percent(pass_rate),
            '内部总人数': internal_total,
            '内部通过人数': internal_pass,
            '内部通过率(%)': _format_percent(internal_pass_rate),
            '渠道总人数': channel_total,
            '渠道通过人数': channel_pass,
            '渠道通过率(%)': _format_percent(channel_pass_rate)
        })

    return pd.DataFrame(result_rows)

def generate_dimension_analysis(all_data: pd.DataFrame, dimension: str) -> pd.DataFrame:
    """
    维度分析。
    区域/分公司：按员工级汇总统计。
    考试科目：按本次实际导入的科目逐科统计，保证产品线表只展示导入科目。
    缺勤口径统一为：该员工所有已导入科目的最终成绩都为0或空。
    """
    if all_data is None or all_data.empty:
        return pd.DataFrame()

    if dimension == '考试科目':
        return _build_subject_dimension_analysis(all_data)

    personal_df = _build_employee_summary(all_data)
    if personal_df.empty or dimension not in personal_df.columns:
        return pd.DataFrame()

    analysis_result = []
    for dim_value, group in personal_df.groupby(dimension, dropna=False):
        dim_value = str(dim_value)
        if dim_value in ['未知区域', '未知分公司', '', 'nan', 'None']:
            continue

        total_people = len(group)
        absent_people = int(group['是否缺勤'].sum())
        attendance_people = total_people - absent_people
        absent_rate = absent_people / total_people * 100 if total_people else 0
        avg_score = group[~group['是否缺勤']]['综合成绩'].mean() if attendance_people else 0.0
        pass_people = len(group[(~group['是否缺勤']) & (group['综合成绩'] >= 80)])
        pass_rate = pass_people / total_people * 100 if total_people else 0

        internal_group = group[group['员工类型'] == '内部']
        internal_total = len(internal_group)
        internal_pass = len(internal_group[(~internal_group['是否缺勤']) & (internal_group['综合成绩'] >= 80)])
        internal_pass_rate = internal_pass / internal_total * 100 if internal_total else 0

        channel_group = group[group['员工类型'] == '渠道']
        channel_total = len(channel_group)
        channel_pass = len(channel_group[(~channel_group['是否缺勤']) & (channel_group['综合成绩'] >= 80)])
        channel_pass_rate = channel_pass / channel_total * 100 if channel_total else 0

        analysis_result.append({
            dimension: dim_value,
            '总人数': total_people,
            '平均分': round(float(avg_score), 2) if pd.notna(avg_score) else 0.0,
            '通过人数': pass_people,
            '通过率(%)': _format_percent(pass_rate),
            '缺勤人数': absent_people,
            '缺勤率(%)': _format_percent(absent_rate),
            '内部总人数': internal_total,
            '内部通过人数': internal_pass,
            '内部通过率(%)': _format_percent(internal_pass_rate),
            '渠道总人数': channel_total,
            '渠道通过人数': channel_pass,
            '渠道通过率(%)': _format_percent(channel_pass_rate)
        })

    result = pd.DataFrame(analysis_result)
    if dimension == '区域' and not result.empty:
        result['_order'] = result[dimension].apply(lambda x: AREA_ORDER.index(x) if x in AREA_ORDER else 99)
        result = result.sort_values('_order').drop(columns=['_order'])
    return result.reset_index(drop=True)


def _get_active_subjects(merged_data: Dict[str, pd.DataFrame]) -> List[str]:
    """按标准科目顺序返回本次实际导入的科目。"""
    if not isinstance(merged_data, dict):
        return []
    return [subject for subject in SUBJECT_LIST if subject in merged_data and isinstance(merged_data.get(subject), pd.DataFrame) and not merged_data.get(subject).empty]


def _display_score_value(value):
    """Excel展示用：未导入考试类型显示/，已导入但空分或0分显示0。"""
    if pd.isna(value):
        return '/'
    try:
        number = float(value)
        if number.is_integer():
            return int(number)
        return round(number, 2)
    except Exception:
        text = str(value).strip()
        return '/' if text in ['', 'nan', 'None', '<NA>'] else text


def generate_dimension_analysis_from_summary(summary_df: pd.DataFrame, active_subjects: List[str], dimension: str) -> pd.DataFrame:
    """
    基于“成绩汇总”sheet的最终平均分生成区域/分公司/产品线分析。

    统一口径：
    1. 平均分 = 取“成绩汇总”sheet中的“平均成绩”列做均值；
    2. 通过人数 = 取“成绩汇总”sheet中的“平均成绩”>=80的人数；
    3. 是否缺勤/缺考 = 取“成绩汇总”sheet中的“是否缺考”列；
    4. 产品线分析按本次导入科目展示，仅统计“成绩汇总”中该科目不是“/”的人员。
    """
    if summary_df is None or summary_df.empty:
        return pd.DataFrame()

    data = summary_df.copy()

    if '工号' not in data.columns:
        return pd.DataFrame()

    if '平均成绩' not in data.columns:
        data['平均成绩'] = 0.0

    data['平均成绩'] = pd.to_numeric(data['平均成绩'], errors='coerce').fillna(0.0)

    if '是否缺考' not in data.columns:
        data['是否缺考'] = data.apply(
            lambda row: all(_is_zero_or_empty_score(row.get(subject, '/')) for subject in active_subjects),
            axis=1
        )
    else:
        data['是否缺考'] = data['是否缺考'].apply(
            lambda x: bool(x) if isinstance(x, bool) else str(x).strip().lower() in ['true', '是', '缺考', '1']
        )

    data['员工类型'] = data['工号'].apply(lambda x: '渠道' if str(x).startswith('S') else '内部')

    def _build_group_row(dim_name: str, dim_value, group: pd.DataFrame, include_absence: bool = True) -> dict:
        total_people = int(len(group))
        absent_people = int(group['是否缺考'].sum()) if '是否缺考' in group.columns else 0
        attendance_people = total_people - absent_people

        # 平均分以“成绩汇总”sheet里的最终平均成绩为准，缺考人员不参与平均分。
        attended_group = group[group['是否缺考'] != True].copy()
        avg_score = attended_group['平均成绩'].mean() if not attended_group.empty else 0.0

        # 通过人数同样以“成绩汇总”sheet里的最终平均成绩为准。
        pass_people = int((attended_group['平均成绩'] >= 80).sum()) if not attended_group.empty else 0
        pass_rate = pass_people / total_people * 100 if total_people else 0

        internal_group = group[group['员工类型'] == '内部'].copy()
        internal_total = int(len(internal_group))
        internal_attended = internal_group[internal_group['是否缺考'] != True].copy()
        internal_pass = int((internal_attended['平均成绩'] >= 80).sum()) if not internal_attended.empty else 0
        internal_pass_rate = internal_pass / internal_total * 100 if internal_total else 0

        channel_group = group[group['员工类型'] == '渠道'].copy()
        channel_total = int(len(channel_group))
        channel_attended = channel_group[channel_group['是否缺考'] != True].copy()
        channel_pass = int((channel_attended['平均成绩'] >= 80).sum()) if not channel_attended.empty else 0
        channel_pass_rate = channel_pass / channel_total * 100 if channel_total else 0

        row = {
            dim_name: dim_value,
            '总人数': total_people,
            '平均分': round(float(avg_score), 2) if pd.notna(avg_score) else 0.0,
            '通过人数': pass_people,
            '通过率(%)': _format_percent(pass_rate),
            '内部总人数': internal_total,
            '内部通过人数': internal_pass,
            '内部通过率(%)': _format_percent(internal_pass_rate),
            '渠道总人数': channel_total,
            '渠道通过人数': channel_pass,
            '渠道通过率(%)': _format_percent(channel_pass_rate)
        }

        if include_absence:
            absent_rate = absent_people / total_people * 100 if total_people else 0
            # 保持区域/分公司分析原有列顺序：缺勤人数、缺勤率放在通过率后。
            row = {
                dim_name: dim_value,
                '总人数': total_people,
                '平均分': round(float(avg_score), 2) if pd.notna(avg_score) else 0.0,
                '通过人数': pass_people,
                '通过率(%)': _format_percent(pass_rate),
                '缺勤人数': absent_people,
                '缺勤率(%)': _format_percent(absent_rate),
                '内部总人数': internal_total,
                '内部通过人数': internal_pass,
                '内部通过率(%)': _format_percent(internal_pass_rate),
                '渠道总人数': channel_total,
                '渠道通过人数': channel_pass,
                '渠道通过率(%)': _format_percent(channel_pass_rate)
            }

        return row

    if dimension == '考试科目':
        rows = []
        for subject in active_subjects:
            if subject not in data.columns:
                continue
            # 产品线分析按“成绩汇总”sheet展示口径：该科目不是“/”才进入该产品线统计。
            subject_group = data[data[subject].apply(lambda x: str(x).strip() not in ['', '/', 'nan', 'None', '<NA>'])].copy()
            if subject_group.empty:
                continue
            rows.append(_build_group_row('考试科目', subject, subject_group, include_absence=False))
        return pd.DataFrame(rows)

    if dimension not in data.columns:
        return pd.DataFrame()

    rows = []
    for dim_value, group in data.groupby(dimension, dropna=False):
        dim_value = str(dim_value)
        if dim_value in ['未知区域', '未知分公司', '', 'nan', 'None']:
            continue
        rows.append(_build_group_row(dimension, dim_value, group, include_absence=True))

    result = pd.DataFrame(rows)
    if dimension == '区域' and not result.empty:
        result['_order'] = result[dimension].apply(lambda x: AREA_ORDER.index(x) if x in AREA_ORDER else 99)
        result = result.sort_values('_order').drop(columns=['_order'])
    return result.reset_index(drop=True)


def generate_summary_excel(merged_data: Dict[str, pd.DataFrame], output_path: str):
    """生成大练兵成绩汇总 Excel。"""
    all_subject_records = []
    output_dir = os.path.dirname(output_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    active_subjects = _get_active_subjects(merged_data)
    drill_season = get_drill_season_from_merged_data(merged_data)

    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for subject in active_subjects:
            df = merged_data.get(subject, pd.DataFrame()) if isinstance(merged_data, dict) else pd.DataFrame()
            if df.empty:
                print(f"⚠️ {subject} 无数据，跳过该sheet")
                continue

            df = df.copy()
            df['最终成绩'] = df.apply(calculate_final_score, axis=1)

            # 单科sheet不再过滤0分或空成绩人员：只要原始导入表中有该人员，就应展示。
            # 单科sheet不展示“是否缺考”；缺考统一在“成绩汇总”sheet中按员工所有科目判断。
            df_for_sheet = df.copy().sort_values('最终成绩', ascending=False)

            sheet_cols = [
                '工号', '姓名', '区域', '分公司', '渠道商名称', '考试科目',
                '第一次考试成绩', '第一次补考成绩', '第二次补考成绩', '最终成绩', '备注'
            ]
            internal_cols = [
                '工号', '姓名', '区域', '分公司', '渠道商名称', '考试科目',
                '第一次考试成绩', '第一次补考成绩', '第二次补考成绩', '最终成绩', '备注'
            ]
            for col in sheet_cols:
                if col not in df_for_sheet.columns:
                    df_for_sheet[col] = ''
            for col in internal_cols:
                if col not in df.columns:
                    df[col] = ''

            df_output = df_for_sheet[sheet_cols].copy()
            for col in ['第一次考试成绩', '第一次补考成绩', '第二次补考成绩', '最终成绩']:
                if col in df_output.columns:
                    df_output[col] = df_output[col].apply(_display_score_value)

            df_output = df_output.fillna('')
            df_output.to_excel(writer, sheet_name=subject, index=False)

            # 后续汇总和维度分析用计算口径数据，不使用展示用“/”。
            all_subject_records.append(df[internal_cols].copy())
            print(f"✅ 已生成科目sheet：{subject}")

        if not all_subject_records:
            pd.DataFrame(columns=['工号', '姓名']).to_excel(writer, sheet_name='空数据', index=False)
            all_data = pd.DataFrame()
        else:
            all_data = pd.concat(all_subject_records, ignore_index=True)

        if not all_data.empty:
            print("\n正在生成成绩汇总Sheet...")
            base_info = all_data[['工号', '区域', '分公司', '渠道商名称', '姓名', '备注']].drop_duplicates(subset='工号', keep='first')

            score_pivot = all_data.pivot_table(index='工号', columns='考试科目', values='最终成绩', aggfunc='first')
            score_pivot = score_pivot.fillna('/')
            for subject in active_subjects:
                if subject not in score_pivot.columns:
                    score_pivot[subject] = '/'
            score_pivot = score_pivot[active_subjects]

            summary_df = base_info.merge(score_pivot, on='工号', how='left')

            def calc_avg(row):
                """
                成绩汇总平均成绩口径：
                1. 某科最终成绩为0、空值或/时，该科不参与统计，并在成绩汇总中展示为/；
                2. 取消尿液、微生物等不同科目的权重计算；
                3. 取该员工所有有效科目最终成绩中的最高分作为“平均成绩”；
                4. 若没有任何有效正分科目，则平均成绩为0。
                """
                valid_scores = []
                for subject in active_subjects:
                    value = row.get(subject, '/')
                    if _is_zero_or_empty_score(value):
                        continue
                    valid_scores.append(_to_numeric_score(value))

                if not valid_scores:
                    return 0.0

                return round(max(valid_scores), 2)

            # 先按原始数值判断是否缺考，再把0/空值展示成/。
            summary_df['是否缺考'] = summary_df.apply(
                lambda row: all(_is_zero_or_empty_score(row.get(subject, '/')) for subject in active_subjects),
                axis=1
            )
            summary_df['平均成绩'] = summary_df.apply(calc_avg, axis=1)

            # 成绩汇总展示口径：某科成绩为0或空时，显示为/，且该科不参与平均成绩计算。
            for subject in active_subjects:
                summary_df[subject] = summary_df[subject].apply(
                    lambda value: '/' if _is_zero_or_empty_score(value) else _display_score_value(value)
                )

            summary_df = summary_df.sort_values('平均成绩', ascending=False)

            final_columns = ['区域', '分公司', '渠道商名称', '工号', '姓名'] + active_subjects + ['平均成绩', '是否缺考', '备注']
            for col in final_columns:
                if col not in summary_df.columns:
                    summary_df[col] = ''
            summary_df[final_columns].to_excel(writer, sheet_name='成绩汇总', index=False)

            try:
                # 区域/分公司/产品线分析统一基于“成绩汇总”sheet中的最终平均分和是否缺考计算，
                # 避免与成绩汇总出现两套平均分或通过人数口径。
                generate_dimension_analysis_from_summary(summary_df[final_columns].copy(), active_subjects, '区域').to_excel(
                    writer, sheet_name='区域分析', index=False
                )
                generate_dimension_analysis_from_summary(summary_df[final_columns].copy(), active_subjects, '分公司').to_excel(
                    writer, sheet_name='分公司分析', index=False
                )
                generate_dimension_analysis_from_summary(summary_df[final_columns].copy(), active_subjects, '考试科目').to_excel(
                    writer, sheet_name='产品线分析', index=False
                )
            except Exception as e:
                print(f"⚠️ 维度分析生成失败：{e}")

        # 异常人员sheet：找不到具体分公司或渠道商时输出，便于排查原始部门字段。
        anomaly_records = []
        if isinstance(merged_data, dict):
            meta = merged_data.get('__meta__', {})
            if isinstance(meta, dict):
                anomaly_records = meta.get('异常人员', []) or []
        if anomaly_records:
            anomaly_df = pd.DataFrame(anomaly_records)
            anomaly_df = anomaly_df.drop_duplicates()
            anomaly_df.to_excel(writer, sheet_name='异常人员', index=False)

        pd.DataFrame(columns=['区域', '应到人数', '实到人数', '缺勤人数', '签到率']).to_excel(writer, sheet_name='区域签到情况', index=False)

        workbook = writer.book
        workbook.properties.title = f'IVD{drill_season}大练兵成绩汇总'
        workbook.properties.subject = f'IVD{drill_season}大练兵'
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            bold_font = Font(bold=True)
            center = Alignment(horizontal='center', vertical='center')
            for cell in sheet[1]:
                cell.font = bold_font
                cell.alignment = center
            for col in sheet.columns:
                max_len = 8
                col_letter = col[0].column_letter
                for cell in col[:300]:
                    if cell.value is not None:
                        max_len = max(max_len, len(str(cell.value)))
                sheet.column_dimensions[col_letter].width = min(max_len + 2, 28)

    print(f"\n🎉 汇总Excel已生成：{output_path}")


def _set_paragraph_text_keep_format(paragraph, text: str):
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _insert_table_after(paragraph, table):
    paragraph._p.addnext(table._tbl)


def _insert_paragraph_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)


def _beautify_docx_table(table, font_size=9):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if table.rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def _df_to_table_image(table_df: pd.DataFrame, img_path: str):
    if table_df is None or table_df.empty:
        table_df = pd.DataFrame({'提示': ['暂无数据']})
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'WenQuanYi Micro Hei']
    plt.rcParams['axes.unicode_minus'] = False
    fig, ax = plt.subplots(figsize=(11, max(len(table_df) * 0.32 + 0.8, 1.2)))
    ax.axis('tight')
    ax.axis('off')
    table = ax.table(cellText=table_df.values, colLabels=table_df.columns, loc='center', cellLoc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(7.5)
    table.scale(1.0, 1.1)
    for (i, j), cell in table.get_celld().items():
        if i == 0:
            cell.set_text_props(weight='bold', fontname='SimHei')
        if j == 0:
            cell.set_text_props(weight='bold', fontname='SimHei')
        cell.set_text_props(fontname='SimHei')
    plt.subplots_adjust(left=0, right=1, top=1, bottom=0)
    plt.savefig(img_path, dpi=300, bbox_inches='tight', pad_inches=0)
    plt.close()


def generate_analysis_report(excel_path: str, template_path: str = '', output_path: str = ''):
    """根据成绩汇总 Excel 和 Word 模板生成大练兵 Word 分析报告。模板为空时自动跳过。"""
    if template_path is None or str(template_path).strip() == '':
        print("ℹ️ 未选择Word模板，跳过Word报告生成，仅生成Excel。")
        return

    print("\n" + "=" * 80)
    print("正在生成Word分析报告...")

    drill_season = get_drill_season_from_path(excel_path, output_path, template_path)
    print(f"✅ Word报告季节名称：{drill_season}大练兵")

    try:
        area_df = pd.read_excel(excel_path, sheet_name='区域分析')
        branch_df = pd.read_excel(excel_path, sheet_name='分公司分析')
        product_df = pd.read_excel(excel_path, sheet_name='产品线分析')
        summary_df = pd.read_excel(excel_path, sheet_name='成绩汇总')
    except Exception as e:
        print(f"❌ 读取Excel失败：{e}")
        return

    try:
        doc = Document(template_path)
        _replace_drill_season_in_doc(doc, drill_season)
    except Exception as e:
        print(f"❌ 打开Word模板失败：{e}")
        return

    total_people = int(area_df['总人数'].sum()) if '总人数' in area_df.columns else len(summary_df)
    total_pass = int(area_df['通过人数'].sum()) if '通过人数' in area_df.columns else int((summary_df.get('平均成绩', 0) >= 80).sum())
    total_pass_rate = round(total_pass / total_people * 100, 2) if total_people else 0
    total_absent = int(area_df['缺勤人数'].sum()) if '缺勤人数' in area_df.columns else 0
    total_absent_rate = round(total_absent / total_people * 100, 2) if total_people else 0
    internal_total = int(area_df['内部总人数'].sum()) if '内部总人数' in area_df.columns else 0
    internal_pass = int(area_df['内部通过人数'].sum()) if '内部通过人数' in area_df.columns else 0
    internal_pass_rate = round(internal_pass / internal_total * 100, 2) if internal_total else 0
    channel_total = int(area_df['渠道总人数'].sum()) if '渠道总人数' in area_df.columns else 0
    channel_pass = int(area_df['渠道通过人数'].sum()) if '渠道通过人数' in area_df.columns else 0
    channel_pass_rate = round(channel_pass / channel_total * 100, 2) if channel_total else 0

    for paragraph in doc.paragraphs:
        if ('本次IVD用户服务工程师' in paragraph.text and '大练兵共有' in paragraph.text) or '本次IVD用户服务工程师春季大练兵共有' in paragraph.text:
            text = (
                f"本次IVD用户服务工程师{drill_season}大练兵共有{total_people}人参加，"
                f"总体通过率（内部工程师80分通过，渠道工程师80分通过）为{total_pass_rate}%，"
                f"缺考率为{total_absent_rate}%。其中内部工程师{internal_total}人，"
                f"通过率为{internal_pass_rate}%；渠道工程师{channel_total}人，"
                f"通过率为{channel_pass_rate}%。总体表现良好。但存在如下突出问题："
            )
            _set_paragraph_text_keep_format(paragraph, text)
            break

    title_paragraphs = {}
    target_titles = [
        '1、区域维度分析',
        '结合通过率以及考勤情况看：',
        '2、分公司维度分析',
        '3、产品线维度分析',
        '4、缺考情况分析',
        '5、各分数段人数占比分析'
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in target_titles:
            title_paragraphs[text] = paragraph

    # 区域维度表
    if '1、区域维度分析' in title_paragraphs and not area_df.empty:
        p = title_paragraphs['1、区域维度分析']
        headers = ['区域', '参考人数', '通过率(%)', '平均成绩', '内部参考人数', '内部通过人数', '内部通过率(%)', '渠道参考人数', '渠道通过人数', '渠道通过率(%)']
        table = doc.add_table(rows=len(area_df) + 1, cols=len(headers))
        table.style = 'Table Grid'
        for j, h in enumerate(headers):
            table.cell(0, j).text = h
        for i, (_, row) in enumerate(area_df.iterrows(), 1):
            values = [row.get('区域', ''), row.get('总人数', ''), row.get('通过率(%)', ''), row.get('平均分', ''), row.get('内部总人数', ''), row.get('内部通过人数', ''), row.get('内部通过率(%)', ''), row.get('渠道总人数', ''), row.get('渠道通过人数', ''), row.get('渠道通过率(%)', '')]
            for j, v in enumerate(values):
                table.cell(i, j).text = str(v)
        _beautify_docx_table(table, 8)
        _insert_table_after(p, table)

    # 考勤表占位
    if '结合通过率以及考勤情况看：' in title_paragraphs:
        p = title_paragraphs['结合通过率以及考勤情况看：']
        headers = ['区域', '考试总人数', '理论签到人数', '签到记录', '签到总人数', '签到率(%)', '考试通过率(%)']
        table = doc.add_table(rows=7, cols=len(headers))
        table.style = 'Table Grid'
        for j, h in enumerate(headers):
            table.cell(0, j).text = h
        for i, region in enumerate(AREA_ORDER, 1):
            table.cell(i, 0).text = region
        _beautify_docx_table(table, 9)
        _insert_table_after(p, table)

    # 分公司/产品线表作为图片插入，防止超宽
    if '2、分公司维度分析' in title_paragraphs:
        filtered = branch_df.drop(columns=['通过人数', '缺勤人数', '缺勤率(%)'], errors='ignore')
        img_path = 'temp_branch_table.png'
        _df_to_table_image(filtered, img_path)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(img_path, width=Inches(6.8))
        if os.path.exists(img_path):
            os.remove(img_path)
        _insert_paragraph_after(title_paragraphs['2、分公司维度分析'], para)

    if '3、产品线维度分析' in title_paragraphs:
        filtered = product_df.drop(columns=['通过人数', '缺勤人数', '缺勤率(%)'], errors='ignore')
        img_path = 'temp_product_table.png'
        _df_to_table_image(filtered, img_path)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(img_path, width=Inches(6.8))
        if os.path.exists(img_path):
            os.remove(img_path)
        _insert_paragraph_after(title_paragraphs['3、产品线维度分析'], para)

    if '4、缺考情况分析' in title_paragraphs and not area_df.empty:
        p = title_paragraphs['4、缺考情况分析']
        headers = ['区域', '应考试人数', '缺考人数', '缺考率(%)']
        table = doc.add_table(rows=len(area_df) + 1, cols=len(headers))
        table.style = 'Table Grid'
        for j, h in enumerate(headers):
            table.cell(0, j).text = h
        for i, (_, row) in enumerate(area_df.iterrows(), 1):
            for j, v in enumerate([row.get('区域', ''), row.get('总人数', ''), row.get('缺勤人数', ''), row.get('缺勤率(%)', '')]):
                table.cell(i, j).text = str(v)
        _beautify_docx_table(table, 9)
        _insert_table_after(p, table)

    if '5、各分数段人数占比分析' in title_paragraphs:
        p = title_paragraphs['5、各分数段人数占比分析']
        bins = [0, 70, 80, 90, 101]
        labels = ['70分以下', '70-80分', '80-90分', '90分以上']
        definitions = ['不及格（重点关注）', '及格（达标人员）', '良好（中坚力量）', '优秀（技术骨干）']
        if '平均成绩' in summary_df.columns:
            counts = pd.cut(summary_df['平均成绩'], bins=bins, labels=labels, right=False).value_counts().reindex(labels).fillna(0).astype(int)
        else:
            counts = pd.Series([0, 0, 0, 0], index=labels)
        ratios = [round(c / total_people * 100, 2) if total_people else 0 for c in counts]
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        for j, h in enumerate(['分数段', '人数', '占比(%)', '定义']):
            table.cell(0, j).text = h
        for i, (seg, count, ratio, definition) in enumerate(zip(labels, counts, ratios, definitions), 1):
            for j, v in enumerate([seg, count, ratio, definition]):
                table.cell(i, j).text = str(v)
        _beautify_docx_table(table, 9)
        _insert_table_after(p, table)

    _replace_drill_season_in_doc(doc, drill_season)

    try:
        out_dir = os.path.dirname(output_path)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)
        doc.save(output_path)
        print(f"\n🎉 Word分析报告生成完成：{os.path.abspath(output_path)}")
    except Exception as e:
        print(f"❌ 保存Word报告失败：{e}")


def main():
    FIRST_FOLDER = r'E:\code\Big-teach\first'
    REMAKE_FOLDER = r'E:\code\Big-teach\remake'
    RESIT_FOLDER = r'E:\code\Big-teach\resit'
    SUMMARY_EXCEL_PATH = get_drill_output_filenames(DEFAULT_DRILL_SEASON)[0]
    TEMPLATE_DOC_PATH = 'IVD用户服务工程师春季大练兵分析报告模版.docx'
    OUTPUT_DOC_PATH = get_drill_output_filenames(DEFAULT_DRILL_SEASON)[1]

    merged_data = merge_all_exam_data(FIRST_FOLDER, REMAKE_FOLDER, RESIT_FOLDER)
    generate_summary_excel(merged_data, SUMMARY_EXCEL_PATH)
    generate_analysis_report(SUMMARY_EXCEL_PATH, TEMPLATE_DOC_PATH, OUTPUT_DOC_PATH)


if __name__ == '__main__':
    main()
